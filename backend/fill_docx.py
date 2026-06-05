import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def set_cell_background(cell, fill_color):
    """Set cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set padding/margins for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Add light grey borders to the table."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # 4/8 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D3D3D3') # light grey
        tblBorders.append(border)
    tblPr.append(tblBorders)

def insert_paragraph_after(para, doc, text="", style='Normal'):
    """Insert a new paragraph after the given paragraph."""
    new_para = doc.add_paragraph(text, style=style)
    para._element.addnext(new_para._element)
    return new_para

def fill_report():
    doc_path = r"d:\jarkom\Laporan Progress_Kelompok 6.docx"
    doc = docx.Document(doc_path)
    
    # Find heading paragraphs
    para_2 = None
    para_3 = None
    para_4 = None
    
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt == 'Gambaran Sistem':
            para_2 = para
        elif txt == 'Progress Pengerjaan Projek':
            para_3 = para
        elif txt == 'Bukti Pengerjaan Projek':
            para_4 = para
            
    if not (para_2 and para_3 and para_4):
        print("Error: Could not find all headings in the document!")
        return
        
    print("Headings found successfully. Inserting text...")
    
    # ---------------------------------------------
    # SECTION 2: Gambaran Sistem
    # ---------------------------------------------
    p = para_2
    p = insert_paragraph_after(p, doc, "Sistem Monitoring Jaringan Lokal Berbasis Web ini dirancang dengan arsitektur terdistribusi untuk mengumpulkan metrik performa jaringan langsung dari perangkat router, memproses data tersebut di backend, menyimpannya di database cloud, dan menampilkannya pada dashboard berbasis web. Sistem ini juga dilengkapi dengan model deteksi anomali berbasis kecerdasan statistik dinamis (Z-Score).")
    
    p = insert_paragraph_after(p, doc, "A. Arsitektur dan Aliran Data Sistem", style='Heading 3')
    p = insert_paragraph_after(p, doc, "Sistem monitoring ini terdiri dari beberapa komponen utama yang saling berintegrasi:")
    p = insert_paragraph_after(p, doc, "1. Data Collector (Backend - Node.js + Puppeteer): Bertindak sebagai mesin pengambil data. Secara terjadwal (setiap 60 detik), backend menjalankan browser headless via Puppeteer untuk masuk ke dashboard admin router ZTE F670L lokal, membuka menu WLAN Status, dan membaca akumulasi jumlah byte diterima (download) dan dikirim (upload). Backend juga mengambil data beban CPU dan uptime sistem operasi Windows lokal melalui PowerShell.")
    p = insert_paragraph_after(p, doc, "2. Database Cloud (Supabase PostgreSQL): Berfungsi sebagai penyimpan data log historis. Data terstruktur dikirimkan dari backend ke tabel network_logs di Supabase Cloud menggunakan Supabase JS client. Kredensial diamankan menggunakan variabel lingkungan (.env).")
    p = insert_paragraph_after(p, doc, "3. Dashboard Pemantau (Frontend - React.js + Recharts + Tailwind CSS): Menyajikan visualisasi data yang interaktif dengan skema Bento Grid modern. Frontend menampilkan kecepatan unduh/unggah aktif, persentase beban CPU lokal, grafik tren lalu lintas data (line/area chart), dan kartu alert anomali. Grafik dirancang secara khusus untuk menandai titik anomali dengan warna merah menyala.")
    p = insert_paragraph_after(p, doc, "4. Tunnelling (Cloudflare Tunnel): Digunakan untuk mengekspos port backend dan frontend lokal ke internet publik melalui domain jarkom.clyuti.my.id. Hal ini memungkinkan dashboard dapat diakses dari mana saja secara real-time tanpa perlu konfigurasi port forwarding pada router ISP.")
    
    p = insert_paragraph_after(p, doc, "B. Mekanisme Pengumpulan Data Web Scraping", style='Heading 3')
    p = insert_paragraph_after(p, doc, "Karena akses statistics router seringkali tidak menyediakan API standar atau SNMP yang terbuka (dikunci oleh ISP), proyek ini mengimplementasikan web scraping menggunakan Puppeteer headless. Alur scraping yang berhasil diterapkan adalah:")
    p = insert_paragraph_after(p, doc, "1. Membuka halaman login router di http://192.168.1.1.")
    p = insert_paragraph_after(p, doc, "2. Memasukkan username ('admin') dan password router (yang tersimpan di .env) pada selector login form (#Frm_Username, #Frm_Password) lalu mengeklik tombol submit (#LoginId).")
    p = insert_paragraph_after(p, doc, "3. Navigasi ke menu utama 'Local Network' (#localnet), dilanjutkan ke sub-menu samping 'Status' (#localNetStatus).")
    p = insert_paragraph_after(p, doc, "4. Klik header bar 'WLAN Status' (#WLANStatusBar) untuk memicu router memuat data bandwidth terbaru menggunakan AJAX.")
    p = insert_paragraph_after(p, doc, "5. Mengekstrak data bytes received dan bytes sent dari ID elemen HTML #TotalBytesCount:0 (misalnya data string '1262835434/3312516106').")
    p = insert_paragraph_after(p, doc, "6. Menghitung kecepatan unduh (Rx) dan unggah (Tx) aktif dalam Mbps dengan membagi selisih byte dari polling sebelumnya terhadap selang waktu (60 detik).")
    
    p = insert_paragraph_after(p, doc, "C. Algoritma Deteksi Anomali Jaringan Berbasis Z-Score", style='Heading 3')
    p = insert_paragraph_after(p, doc, "Untuk memberikan nilai tambah bagi sistem pemantauan, kami mengimplementasikan deteksi anomali lalu lintas data dan beban CPU berbasis statistik:")
    p = insert_paragraph_after(p, doc, "- Proses Training Dinamis: Setiap 5 menit sekali, backend mengambil 100 data log terbaru dari Supabase untuk menghitung nilai Rata-rata (mean/μ) dan Standar Deviasi (stdDev/σ) dari metrik cpu_usage, rx_speed_mbps, dan tx_speed_mbps.")
    p = insert_paragraph_after(p, doc, "- Perhitungan Z-Score: Setiap data baru masuk, deviasi nilainya diukur dengan rumus: Z = |x - μ| / σ. Jika skor Z melampaui ambang batas sensitivitas (ANOMALY_THRESHOLD = 3.0) dan metrik melampaui ambang batas absolut minimum (misalnya CPU > 60%, Rx > 50 Mbps, Tx > 20 Mbps), sistem menandai data tersebut sebagai anomali.")
    p = insert_paragraph_after(p, doc, "- Mekanisme Cadangan (Cold Start): Jika database memiliki kurang dari 20 record (sehingga standar deviasi belum stabil), sistem secara otomatis beralih ke batas statis (CPU > 85%, Rx > 150 Mbps, Tx > 50 Mbps) sebagai pengaman sementara.")
    
    # ---------------------------------------------
    # SECTION 3: Progress Pengerjaan Projek
    # ---------------------------------------------
    p3 = para_3
    p3 = insert_paragraph_after(p3, doc, "Proyek pengembangan dashboard monitoring jaringan ini telah diselesaikan hingga tahap akhir (fungsionalitas penuh). Seluruh komponen backend, frontend, database cloud, dan hardware integration telah berhasil diimplementasikan dan diuji. Berikut adalah ringkasan progress pengerjaan proyek Kelompok 6:")
    
    # We will insert a table after p3. To do that, we can add it to the document and move its element.
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tahap Pengerjaan'
    hdr_cells[1].text = 'Fitur / Detail Implementasi'
    hdr_cells[2].text = 'Status'
    for cell in hdr_cells:
        set_cell_background(cell, '1F4E78') # dark blue header
        set_cell_margins(cell)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    tasks = [
        ("Tahap 1: Cloud & Database Setup", "Inisialisasi database Supabase Cloud, pembuatan tabel 'network_logs', konfigurasi skema tabel, dan penyesuaian kebijakan RLS (Row-Level Security) agar database dapat diakses backend secara aman.", "100% (Selesai)"),
        ("Tahap 2: Backend Development", "Pengambilan metrik performa komputer lokal (Uptime & CPU usage) menggunakan perintah PowerShell Windows di Node.js. Koneksi API Supabase client untuk sinkronisasi data log.", "100% (Selesai)"),
        ("Tahap 3: Model Deteksi Anomali", "Pengembangan algoritma Z-Score multi-dimensi secara dinamis dengan retraining berkala setiap 5 menit menggunakan data historis Supabase. Implementasi fallback batas statis.", "100% (Selesai)"),
        ("Tahap 4: Frontend Bento Grid UI", "Perancangan dan pembangunan antarmuka dashboard modern menggunakan React.js, Tailwind CSS, Lucide Icons, dan visualisasi grafik real-time dari Recharts yang mendukung highlight anomali (titik merah).", "100% (Selesai)"),
        ("Tahap 5: Refactoring Data Collector", "Migrasi pengambil metrik jaringan dari local adapter WiFi laptop ke scraping langsung router ZTE F670L (menggunakan Puppeteer headless browser untuk login, navigasi menu Status, ekspansi WLAN Status, dan parsing AJAX bytes).", "100% (Selesai)"),
        ("Tahap 6: Tunnelling & Mock Mode", "Konfigurasi Cloudflare Tunnel (cloudflared) untuk pemetaan port lokal ke domain publik jarkom.clyuti.my.id. Penyediaan fitur Mock Mode di backend untuk memicu skenario anomali secara terencana saat presentasi kelas.", "100% (Selesai)")
    ]
    
    for t_name, t_desc, t_status in tasks:
        row_cells = table.add_row().cells
        row_cells[0].text = t_name
        row_cells[1].text = t_desc
        row_cells[2].text = t_status
        for i, cell in enumerate(row_cells):
            set_cell_margins(cell)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
            elif i == 2:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(34, 139, 34) # Forest Green
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    # Move table to correct position in document (after p3)
    p3._element.addnext(table._element)
    
    # Add a spacer paragraph after table
    p_after_tbl = doc.add_paragraph("")
    table._element.addnext(p_after_tbl._element)
    
    # ---------------------------------------------
    # SECTION 4: Bukti Pengerjaan Projek
    # ---------------------------------------------
    p4 = para_4
    p4 = insert_paragraph_after(p4, doc, "Sebagai bukti pengerjaan proyek, berikut adalah tangkapan layar (screenshot) dari sistem monitoring jaringan yang telah berhasil diimplementasikan:")
    
    # We have some images in the brain folder:
    # C:\Users\Asus\.gemini\antigravity-ide\brain\aac3ad52-66a8-406b-9642-8b44dbfea335\media__1779437318516.png (Main dashboard normal)
    # C:\Users\Asus\.gemini\antigravity-ide\brain\aac3ad52-66a8-406b-9642-8b44dbfea335\media__1779438624030.png (Main dashboard anomaly)
    # C:\Users\Asus\.gemini\antigravity-ide\brain\aac3ad52-66a8-406b-9642-8b44dbfea335\media__1779437524399.png (Supabase dashboard logs)
    # d:\jarkom\backend\router_debug.png (Router status scraping debug screenshot)
    
    images = [
        (r"C:\Users\Asus\.gemini\antigravity-ide\brain\aac3ad52-66a8-406b-9642-8b44dbfea335\media__1779437318516.png", "Gambar 1: Tampilan Utama Dashboard Monitoring Jaringan Berbasis Bento Grid (Keadaan Normal)"),
        (r"C:\Users\Asus\.gemini\antigravity-ide\brain\aac3ad52-66a8-406b-9642-8b44dbfea335\media__1779438624030.png", "Gambar 2: Tampilan Dashboard Saat Mendeteksi Anomali Jaringan (Kotak Alert Merah Berkedip)"),
        (r"C:\Users\Asus\.gemini\antigravity-ide\brain\aac3ad52-66a8-406b-9642-8b44dbfea335\media__1779437524399.png", "Gambar 3: Tabel Riwayat Log Trafik di Supabase Cloud Database"),
        (r"d:\jarkom\backend\router_debug.png", "Gambar 4: Halaman Manajemen ZTE Router F670L yang Di-scrape oleh Puppeteer")
    ]
    
    for img_path, caption in images:
        if os.path.exists(img_path):
            # Insert image paragraph
            p_img = insert_paragraph_after(p4, doc, "")
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_img.add_run()
            try:
                run.add_picture(img_path, width=Inches(5.5))
                # Insert caption paragraph
                p_cap = insert_paragraph_after(p_img, doc, caption, style='Normal')
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.runs[0].font.italic = True
                p_cap.runs[0].font.size = Pt(9.5)
                p4 = p_cap
                print(f"Inserted image: {img_path}")
            except Exception as e:
                print(f"Error inserting image {img_path}: {e}")
        else:
            print(f"Warning: Image file not found: {img_path}")
            
    try:
        doc.save(doc_path)
        print(f"Document saved successfully to: {doc_path}")
    except PermissionError:
        alt_path = r"d:\jarkom\Laporan Progress_Kelompok 6_Filled.docx"
        doc.save(alt_path)
        print(f"Warning: Primary file is locked by another process (likely open in Word).")
        print(f"Successfully saved to alternative path: {alt_path}")

if __name__ == "__main__":
    fill_report()
