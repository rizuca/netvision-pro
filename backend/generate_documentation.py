import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def set_cell_background(cell, hex_color):
    """Set the background color of a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner margins (padding) of a table cell in twentieths of a point (dxa)."""
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)

def set_table_borders(table):
    """Add clean light-gray borders to a table."""
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # thin border
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'CBD5E1')  # gray-300
            borders.append(border)
        tbl_pr[0].append(borders)

def format_run(run, font_name="Arial", size_pt=11, color_rgb=(30, 41, 59), bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.bold = bold
    run.italic = italic

def add_styled_heading(doc, text, level, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    if level == 1:
        format_run(run, font_name="Arial", size_pt=18, color_rgb=(15, 23, 42), bold=True)
    elif level == 2:
        format_run(run, font_name="Arial", size_pt=14, color_rgb=(37, 99, 235), bold=True)
    elif level == 3:
        format_run(run, font_name="Arial", size_pt=12, color_rgb=(71, 85, 105), bold=True)
    return p

def create_document():
    doc = docx.Document()
    
    # Page setup - Standard margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Header Page
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(50)
    p_title.paragraph_format.space_after = Pt(10)
    run_title = p_title.add_run("DOKUMENTASI TEKNIS & MANUAL PANDUAN")
    format_run(run_title, font_name="Arial", size_pt=24, color_rgb=(15, 23, 42), bold=True)
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_subtitle.paragraph_format.space_after = Pt(40)
    run_sub = p_subtitle.add_run("NetVision Pro: Network Monitoring Dashboard & Dynamic Anomaly Detection\nJaringan WiFi Router ZTE F670L")
    format_run(run_sub, font_name="Arial", size_pt=14, color_rgb=(71, 85, 105), bold=False, italic=True)
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(100)
    run_meta = p_meta.add_run("Kelompok 6 - Jaringan Komputer\nUniversitas Negeri / Kampus\nTahun 2026")
    format_run(run_meta, font_name="Arial", size_pt=11, color_rgb=(100, 116, 139), bold=True)
    
    doc.add_page_break()

    # BAB I: PENDAHULUAN
    add_styled_heading(doc, "BAB I: PENDAHULUAN", level=1)
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1.15
    p1.paragraph_format.space_after = Pt(8)
    run = p1.add_run(
        "NetVision Pro adalah sebuah aplikasi monitoring jaringan nirkabel (wireless network) modern "
        "yang dirancang untuk mengawasi lalu lintas data (bandwidth), penggunaan sumber daya server (CPU Load), "
        "dan stabilitas koneksi secara real-time. Dilengkapi dengan mesin kecerdasan statistik berbasis "
        "algoritma Z-Score dinamis, aplikasi ini dapat mendeteksi adanya keanehan perilaku jaringan (anomali) "
        "seperti lonjakan trafik mendadak (traffic spikes) atau beban prosesor berlebih (CPU overload) "
        "yang berpotensi mengganggu performa jaringan nirkabel."
    )
    format_run(run)
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1.15
    p2.paragraph_format.space_after = Pt(8)
    run = p2.add_run(
        "Berbeda dengan aplikasi pemantau jaringan konvensional yang mengandalkan batasan statis (fixed thresholds) "
        "yang seringkali memicu alarm palsu (false alarms), NetVision Pro melatih model anomalinya secara dinamis "
        "setiap 5 menit sekali dengan mengambil 100 riwayat rekaman data sebelumnya dari Supabase Cloud Database. "
        "Aplikasi ini secara khusus diintegrasikan dengan Router ZTE F670L (IP default 192.168.1.1) untuk membaca "
        "lalu lintas byte jaringan dan mendata setiap klien yang terhubung ke titik akses WiFi (WLAN)."
    )
    format_run(run)

    # BAB II: ARSITEKTUR SISTEM & STRUKTUR KODE
    add_styled_heading(doc, "BAB II: ARSITEKTUR SISTEM & STRUKTUR PROGRAM", level=1)
    
    p3 = doc.add_paragraph()
    p3.paragraph_format.line_spacing = 1.15
    p3.paragraph_format.space_after = Pt(8)
    run = p3.add_run(
        "NetVision Pro dibangun menggunakan arsitektur client-server full-stack dengan integrasi database "
        "awan (Cloud Database). Berikut adalah tiga komponen utama penyusun arsitektur sistem:"
    )
    format_run(run)
    
    # List of components
    bullet_items = [
        ("Backend (Node.js & Express): ", "Bertanggung jawab atas penjadwalan berkala (polling) pengambilan data sistem lokal dan router melalui Puppeteer Headless. Komponen ini juga memproses model Z-score untuk mendeteksi anomali sebelum menyimpan data ke database."),
        ("Frontend (React.js & Vite): ", "Berfungsi sebagai antarmuka visual (dashboard) pengguna dengan tata letak Bento Grid modern, grafik lalu lintas interaktif menggunakan pustaka Recharts, lencana indikator status live, dan panel informasi teknis."),
        ("Supabase Cloud Database: ", "Berperan sebagai media penyimpanan historis logs. RLS (Row-Level Security) dinonaktifkan pada tabel 'network_logs' untuk mengizinkan backend menulis entri data secara lancar."),
        ("Cloudflare Tunnel (Tunneling): ", "Digunakan untuk mengubungkan localhost port 5000 (backend + static frontend) ke domain publik (jarkom.clyuti.my.id) secara aman tanpa perlu melakukan port-forwarding pada router rumah.")
    ]
    
    for title, desc in bullet_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r_title = p.add_run(title)
        format_run(r_title, bold=True)
        r_desc = p.add_run(desc)
        format_run(r_desc)

    p_struct = doc.add_paragraph()
    p_struct.paragraph_format.space_before = Pt(8)
    p_struct.paragraph_format.space_after = Pt(6)
    run = p_struct.add_run("Struktur Direktori Proyek:")
    format_run(run, bold=True)
    
    # Table for directory structure
    dir_table = doc.add_table(rows=1, cols=3)
    dir_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(dir_table)
    
    hdr_cells = dir_table.rows[0].cells
    hdr_cells[0].text = 'Direktori / File'
    hdr_cells[1].text = 'Teknologi / Library'
    hdr_cells[2].text = 'Deskripsi / Peran Utama'
    for cell in hdr_cells:
        set_cell_background(cell, 'F1F5F9')
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        format_run(cell.paragraphs[0].runs[0], bold=True, color_rgb=(15, 23, 42))

    directories_data = [
        ('backend/server.js', 'Express, Puppeteer, SupabaseJS', 'Titik masuk utama backend. Mengatur siklus poll, scraping data router, perhitungan deteksi anomali Z-score, dan menyediakan REST API.'),
        ('backend/.env', 'Dotenv', 'Penyimpan konfigurasi variabel lingkungan seperti ROUTER_PASS, SUPABASE_URL, MOCK_MODE, dan PORT.'),
        ('frontend/src/components/Dashboard.jsx', 'React, Recharts, Lucide Icons', 'Komponen utama dasbor antarmuka pengguna (UI). Berisi grid bento, pie chart CPU load, grafik garis lalu lintas data, serta pemanggil API.'),
        ('frontend/src/index.css', 'TailwindCSS, Vanilla CSS', 'Pengatur tema styling visual, font (Plus Jakarta Sans), animasi micro-interactions, dan layout mode terang (light mode).'),
        ('frontend/dist/', 'HTML / JS / CSS Kompilasi', 'Folder aset statis hasil kompilasi (Vite build) yang didistribusikan secara langsung oleh Express backend.')
    ]

    for path, tech, description in directories_data:
        row_cells = dir_table.add_row().cells
        row_cells[0].text = path
        row_cells[1].text = tech
        row_cells[2].text = description
        for cell in row_cells:
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            format_run(cell.paragraphs[0].runs[0], size_pt=9.5)

    doc.add_page_break()

    # BAB III: MEKANISME PENGUMPULAN DATA (DATA COLLECTION)
    add_styled_heading(doc, "BAB III: MEKANISME PENGUMPULAN DATA (DATA COLLECTION)", level=1)
    
    p4 = doc.add_paragraph()
    p4.paragraph_format.line_spacing = 1.15
    p4.paragraph_format.space_after = Pt(8)
    run = p4.add_run(
        "Sistem monitoring NetVision Pro beroperasi secara mandiri (autonomous) dengan melakukan jajak pendapat "
        "(polling) berkala setiap 60 detik. Berikut adalah rincian asal-usul pengambilan data metrik:"
    )
    format_run(run)
    
    # 3.1 CPU dan Uptime
    add_styled_heading(doc, "3.1 Penggunaan CPU dan Uptime Sistem (Windows)", level=2)
    p_cpu = doc.add_paragraph()
    p_cpu.paragraph_format.line_spacing = 1.15
    p_cpu.paragraph_format.space_after = Pt(8)
    run = p_cpu.add_run(
        "Metrik performa komputer server diperoleh dengan memanfaatkan utilitas baris perintah PowerShell Windows. "
        "Perintah dijalankan secara sinkron di latar belakang menggunakan pustaka child_process (execSync) bawaan Node.js:\n"
        "1. CPU Load (%): Didapatkan dari kelas WMI Win32_Processor dengan perintah:\n"
        "   Get-WimObject Win32_Processor | Measure-Object -Property LoadPercentage -Average\n"
        "2. System Uptime: Dihitung dengan mengambil waktu booting sistem terakhir dari kelas Win32_OperatingSystem "
        "lalu dikurangkan dengan waktu sekarang, menghasilkan durasi aktif komputer dalam format jam, menit, detik."
    )
    format_run(run)

    # 3.2 Rx/Tx Bytes
    add_styled_heading(doc, "3.2 Data Lalu Lintas (Rx/Tx) dari Router ZTE F670L", level=2)
    p_rxtx = doc.add_paragraph()
    p_rxtx.paragraph_format.line_spacing = 1.15
    p_rxtx.paragraph_format.space_after = Pt(8)
    run = p_rxtx.add_run(
        "Karena router nirkabel tidak menyediakan protokol SNMP secara terbuka, NetVision Pro menggunakan teknik "
        "Web Scraping headless dengan Puppeteer:\n"
        "1. Otentikasi: Browser Chrome tak kasat mata (headless) dijalankan, membuka halaman login http://192.168.1.1, "
        "memasukkan username (admin) dan password router (ROUTER_PASS), kemudian menekan tombol login.\n"
        "2. Navigasi Menu: Puppeteer melakukan klik berurutan pada menu 'Local Network' -> 'Status' -> collapsible header 'WLAN Status'.\n"
        "3. Pengambilan Nilai Byte: Menunggu permintaan AJAX internal selesai, kemudian membaca atribut 'title' pada elemen "
        "<span id=\"TotalBytesCount:0\">. Elemen ini menyimpan total byte terkirim dan diterima router sejak menyala dengan format 'RxBytes/TxBytes' (misal: 478922237/627385569).\n"
        "4. Konversi Kecepatan (Mbps): Karena data di router berbentuk akumulasi byte, backend menghitung selisih (delta) "
        "data saat ini dengan data pada 60 detik lalu (siklus sebelumnya), kemudian dikonversi menjadi megabit per detik:\n"
        "   Kecepatan (Mbps) = (Byte Baru - Byte Lama) * 8 / (Durasi Detik * 1.000.000)"
    )
    format_run(run)

    # 3.3 Clients
    add_styled_heading(doc, "3.3 Daftar Klien WiFi (WLAN Client Status)", level=2)
    p_clients = doc.add_paragraph()
    p_clients.paragraph_format.line_spacing = 1.15
    p_clients.paragraph_format.space_after = Pt(8)
    run = p_clients.add_run(
        "Selain lalu lintas byte, daftar klien nirkabel yang sedang tersambung ke WiFi dikumpulkan dari menu "
        "'WLAN Client Status' di router. Puppeteer mengekspansi bar status perangkat nirkabel, lalu memindai struktur HTML "
        "kontainer '#Wlan_ClientStat_container'. Untuk setiap baris kartu kontainer klien yang aktif, sistem menyaring:\n"
        "- HostName: Nama perangkat klien (contoh: ASUS-ROG-Laptop, iPhone-15-Pro, OPPO-A92).\n"
        "- IPAddress: Alamat IP lokal klien yang diberikan oleh server DHCP router.\n"
        "- MACAddress: Alamat fisik unik kartu jaringan klien.\n"
        "- ESSID: SSID WiFi tempat klien terhubung (contoh: Winaila 5G)."
    )
    format_run(run)

    doc.add_page_break()

    # BAB IV: DETEKSI ANOMALI BERBASIS Z-SCORE
    add_styled_heading(doc, "BAB IV: ALGORITMA DETEKSI ANOMALI (Z-SCORE)", level=1)
    
    p_z1 = doc.add_paragraph()
    p_z1.paragraph_format.line_spacing = 1.15
    p_z1.paragraph_format.space_after = Pt(8)
    run = p_z1.add_run(
        "Deteksi anomali pada NetVision Pro mengimplementasikan teknik statistik Z-Score Multi-Dimensi. "
        "Z-score mengukur seberapa jauh suatu nilai data menyimpang dari rata-rata (mean) historis kelompok datanya, "
        "dinyatakan dalam satuan standar deviasi. Rumus matematika Z-score didefinisikan sebagai:"
    )
    format_run(run)
    
    p_formula = doc.add_paragraph()
    p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_formula.paragraph_format.space_before = Pt(8)
    p_formula.paragraph_format.space_after = Pt(8)
    run_f = p_formula.add_run("Z = | X - μ | / σ")
    format_run(run_f, font_name="Consolas", size_pt=14, color_rgb=(37, 99, 235), bold=True)
    
    p_formula_desc = doc.add_paragraph()
    p_formula_desc.paragraph_format.line_spacing = 1.15
    p_formula_desc.paragraph_format.space_after = Pt(8)
    run = p_formula_desc.add_run(
        "Di mana:\n"
        "- X: Nilai metrik real-time terbaru yang sedang diuji.\n"
        "- μ (Mu): Nilai rata-rata (mean) dari 100 logs historis Supabase.\n"
        "- σ (Sigma): Standar deviasi dari 100 logs historis Supabase."
    )
    format_run(run, italic=True)

    add_styled_heading(doc, "4.1 Pelatihan Ulang Dinamis (Dynamic Re-training)", level=2)
    p_train = doc.add_paragraph()
    p_train.paragraph_format.line_spacing = 1.15
    p_train.paragraph_format.space_after = Pt(8)
    run = p_train.add_run(
        "Model statistik dilatih ulang (re-trained) di latar belakang secara otomatis setiap 5 menit. "
        "Perubahan pola jaringan harian (seperti sepi di malam hari dan ramai di siang hari) akan diadaptasi oleh nilai "
        "mean (μ) dan standar deviasi (σ) yang terus diperbarui. Ini mencegah alarm palsu akibat perubahan pola alami jaringan."
    )
    format_run(run)

    add_styled_heading(doc, "4.2 Batas Pemicu Anomali (Anomaly Threshold)", level=2)
    p_trigger = doc.add_paragraph()
    p_trigger.paragraph_format.line_spacing = 1.15
    p_trigger.paragraph_format.space_after = Pt(8)
    run = p_trigger.add_run(
        "Sistem menandai adanya pencilan/anomali (is_anomaly = true) jika nilai Z-score suatu metrik melebihi batas "
        "ANOMALY_THRESHOLD (default: 3.0 standar deviasi). Untuk mencegah alarm palsu pada trafik rendah, sistem juga menyaratkan "
        "batas bawah nilai fisik minimum:\n"
        "1. CPU Load: Z-score > 3.0 dan beban CPU harus di atas 60%.\n"
        "2. Download (Rx): Z-score > 3.0 dan kecepatan harus di atas 50 Mbps.\n"
        "3. Upload (Tx): Z-score > 3.0 dan kecepatan harus di atas 20 Mbps.\n\n"
        "Jika salah satu atau lebih metrik memenuhi kondisi tersebut, status anomali diaktifkan "
        "dan alasan anomali (anomaly_reason) akan digabungkan (misal: 'CPU Load abnormal: 96% & Trafik Download abnormal: 72 Mbps')."
    )
    format_run(run)

    doc.add_page_break()

    # BAB V: PANDUAN ANTARMUKA PENGGUNA (DASHBOARD WIDGETS)
    add_styled_heading(doc, "BAB V: PANDUAN ANTARMUKA PENGGUNA (DASHBOARD WIDGETS)", level=1)
    
    p5 = doc.add_paragraph()
    p5.paragraph_format.line_spacing = 1.15
    p5.paragraph_format.space_after = Pt(8)
    run = p5.add_run(
        "Antarmuka dasbor NetVision Pro didesain dengan konsep Bento Grid bergaya modern yang bersih "
        "menggunakan tema mode terang (light mode). Berikut adalah fungsi dan mekanisme kerja masing-masing widget:"
    )
    format_run(run)

    widgets = [
        ("1. Header Utama (Sistem Bar): ", 
         "Terletak di bagian paling atas dasbor. Berisi nama aplikasi 'NetVision Pro', target IP router (192.168.1.1), "
         "indikator status koneksi backend (🟢 LIVE atau 🟣 MOCK MODE), serta lencana lari real-time jumlah klien WiFi yang tersambung (👥 X Client). "
         "Dilengkapi tombol manual refresh (ikon putar) untuk mempercepat pembaruan data sebelum siklus 60 detik."),
        
        ("2. Widget Download Speed (Rx): ", 
         "Menampilkan kecepatan unduh real-time dalam satuan Mbps dengan angka berukuran ekstra besar di pojok kiri atas grid. "
         "Dilengkapi progress bar horizontal yang bergerak menyesuaikan beban kecepatan dari skala 0-100 Mbps, serta catatan kaki "
         "berupa akumulasi data terunduh total (Total RX) dalam Gigabyte (GB)."),
        
        ("3. Widget Upload Speed (Tx): ", 
         "Menampilkan kecepatan unggah real-time saat ini. Bergerak dengan skala visual progress bar 0-50 Mbps, "
         "dan mencatat total volume data yang telah diunggah (Total TX) oleh router dalam satuan Gigabyte (GB)."),
        
        ("4. Widget System Uptime: ", 
         "Menampilkan durasi aktifnya sistem operasi komputer server dalam format Jam:Menit:Detik (contoh: 125h 38m 36s). "
         "Widget ini berfungsi untuk memastikan server berjalan stabil tanpa reboot yang tidak direncanakan. Siklus polling "
         "lokal dicatat di bagian bawah widget ini."),
        
        ("5. Widget CPU Load (Gauge Chart): ", 
         "Menampilkan grafik lingkaran setengah (PieChart) yang dinamis visualnya. Membagi area CPU terpakai (biru) dan bebas (abu-abu). "
         "Di bagian tengah grafik lingkaran, persentase CPU ditampilkan dalam teks tebal. Widget ini secara otomatis mendeteksi "
         "jika beban CPU melebihi 75% dan melabelinya sebagai 'Beban Berat' (teks merah) atau 'Normal' (teks hijau)."),
        
        ("6. Grafik Utama Live Traffic (Bandwidth): ", 
         "Sebuah LineChart interaktif berukuran besar (col-span 3) yang memvisualisasikan lalu lintas data historis. "
         "Garis biru melambangkan kecepatan download dan garis fuchsia melambangkan upload. Jika data pada waktu tertentu "
         "dinyatakan sebagai anomali oleh model Z-Score, grafik akan secara otomatis memunculkan titik merah menyala (glow red dot) "
         "sebagai penanda visual lonjakan berbahaya."),
        
        ("7. Widget Status Keamanan Jaringan: ", 
         "Sebuah panel visual yang memancarkan status keamanan saat ini. Jika kondisi normal, panel berwarna hijau (emerald) "
         "dengan label 'Status: AMAN'. Jika anomali dipicu, panel berkedip merah terang dengan label 'ANOMALI TERDETEKSI' "
         "dan membeberkan kalimat alasan anomali."),
        
        ("8. Panel Klien Terkoneksi (WLAN Clients): ", 
         "Menampilkan daftar seluruh gawai (devices) yang terhubung ke jaringan nirkabel secara detail. Widget ini memuat "
         "Nama Host perangkat (seperti OPPO-A92, ASUS-ROG), alamat IP lokal, MAC Address fisik gawai, dan SSID yang terhubung. "
         "Membantu administrator jaringan memantau jika ada perangkat asing mencurigakan masuk."),
        
        ("9. Panel Log Database Historis (Supabase): ", 
         "Menampilkan 10 logs data terakhir yang tersimpan di awan Supabase dalam bentuk tabel interaktif. Memuat kolom waktu, "
         "beban CPU, kecepatan Rx (Mbps), kecepatan Tx (Mbps), serta indikator badge status apakah data tersebut bertipe NORMAL (hijau) "
         "atau ANOMALI (merah).")
    ]

    for title, desc in widgets:
        p_w = doc.add_paragraph()
        p_w.paragraph_format.line_spacing = 1.15
        p_w.paragraph_format.space_after = Pt(6)
        p_w.paragraph_format.left_indent = Inches(0.25)
        r_title = p_w.add_run(title)
        format_run(r_title, bold=True, color_rgb=(37, 99, 235))
        r_desc = p_w.add_run(desc)
        format_run(r_desc)

    doc.add_page_break()

    # BAB VI: KESIMPULAN
    add_styled_heading(doc, "BAB VI: KESIMPULAN", level=1)
    
    p6 = doc.add_paragraph()
    p6.paragraph_format.line_spacing = 1.15
    p6.paragraph_format.space_after = Pt(8)
    run = p6.add_run(
        "NetVision Pro merupakan perpaduan antara teknologi web scraping Puppeteer, scripting sistem operasi lokal, "
        "analisis data statistik real-time, dan visualisasi data dasbor modern. Dengan mengandalkan model statistik Z-Score "
        "yang dinamis dan adaptif terhadap pola harian, administrator jaringan dapat mengidentifikasi masalah performa WiFi "
        "secara cepat, tepat, dan otomatis."
    )
    format_run(run)

    # Save document
    output_path = r"d:\jarkom\Dokumentasi_Aplikasi_NetVision_Pro.docx"
    doc.save(output_path)
    print(f"Sukses! Dokumentasi dibuat di {output_path}")

if __name__ == "__main__":
    create_document()
