# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime

OUTPUT_PATH = r"d:\jarkom\Dokumentasi_Lengkap_NetVision_Pro.docx"

# ─────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc.append(shd)

def set_cell_padding(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tc.append(tcMar)

def add_table_borders(table, color='CBD5E1'):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tblPr.append(borders)

def set_para_spacing(para, before=0, after=6, line=1.15):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line

def add_run(para, text, bold=False, italic=False, size=11,
            color=(30, 41, 59), font='Arial'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    return run

def heading(doc, text, level=1, before=14, after=6):
    """Add a styled heading paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    if level == 1:
        add_run(p, text, bold=True, size=17, color=(15,23,42))
    elif level == 2:
        add_run(p, text, bold=True, size=13, color=(37,99,235))
    elif level == 3:
        add_run(p, text, bold=True, size=11, color=(71,85,105))
    return p

def body(doc, text, before=0, after=8, italic=False):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    set_para_spacing(p, before, after)
    add_run(p, text, italic=italic)
    return p

def bullet(doc, label, desc, before=0, after=5):
    p = doc.add_paragraph(style='List Bullet')
    set_para_spacing(p, before, after)
    if label:
        add_run(p, label, bold=True, color=(37,99,235))
    add_run(p, desc)
    return p

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('─' * 90)
    run.font.color.rgb = RGBColor(203, 213, 225)
    run.font.size = Pt(8)
    return p

def code_block(doc, lines):
    """Simulate a code block with a shaded table."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, '1E293B')
    set_cell_padding(cell, top=140, bottom=140, left=200, right=200)
    cell.paragraphs[0].clear()
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        r = p.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(186, 230, 253)  # sky-200
    return t

def two_col_table(doc, rows_data, headers=None, header_bg='1E3A5F'):
    """Create a 2-column info table."""
    total_rows = len(rows_data) + (1 if headers else 0)
    t = doc.add_table(rows=total_rows, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_table_borders(t)
    row_idx = 0
    if headers:
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            set_cell_bg(hdr[i], header_bg)
            set_cell_padding(hdr[i], 120, 120, 160, 160)
            r = hdr[i].paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(255,255,255)
        row_idx = 1
    for label, val in rows_data:
        cells = t.rows[row_idx].cells
        set_cell_bg(cells[0], 'F8FAFC')
        set_cell_padding(cells[0], 100, 100, 160, 160)
        set_cell_padding(cells[1], 100, 100, 160, 160)
        r0 = cells[0].paragraphs[0].add_run(label)
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = RGBColor(51,65,85)
        r1 = cells[1].paragraphs[0].add_run(val)
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(30,41,59)
        row_idx += 1
    return t

def multi_col_table(doc, col_headers, rows_data, header_bg='1E3A5F',
                    col_widths=None):
    total_rows = len(rows_data) + 1
    t = doc.add_table(rows=total_rows, cols=len(col_headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_table_borders(t)
    # Header
    hcells = t.rows[0].cells
    for i, h in enumerate(col_headers):
        set_cell_bg(hcells[i], header_bg)
        set_cell_padding(hcells[i], 120, 120, 150, 150)
        r = hcells[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255,255,255)
    # Rows
    for ri, row_vals in enumerate(rows_data):
        rcells = t.rows[ri+1].cells
        for ci, val in enumerate(row_vals):
            set_cell_padding(rcells[ci], 100, 100, 150, 150)
            rr = rcells[ci].paragraphs[0].add_run(str(val))
            rr.font.size = Pt(10)
            rr.font.color.rgb = RGBColor(30,41,59)
            if ci == 0:
                rr.bold = True
    return t

# ─────────────────────────────────────────────
# MAIN DOCUMENT BUILDER
# ─────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.1)
        sec.right_margin  = Inches(1.0)

    # ═══════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_title.paragraph_format.space_before = Pt(60)
    cover_title.paragraph_format.space_after  = Pt(6)
    r = cover_title.add_run('DOKUMENTASI TEKNIS LENGKAP')
    r.font.name = 'Arial'; r.font.size = Pt(26); r.bold = True
    r.font.color.rgb = RGBColor(15,23,42)

    cover_sub = doc.add_paragraph()
    cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_sub.paragraph_format.space_after = Pt(4)
    r = cover_sub.add_run('NetVision Pro')
    r.font.name = 'Arial'; r.font.size = Pt(20); r.bold = True
    r.font.color.rgb = RGBColor(37,99,235)

    cover_sub2 = doc.add_paragraph()
    cover_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_sub2.paragraph_format.space_after = Pt(40)
    r = cover_sub2.add_run('Network Monitoring Dashboard & Dynamic Anomaly Detection System')
    r.font.name = 'Arial'; r.font.size = Pt(13); r.italic = True
    r.font.color.rgb = RGBColor(71,85,105)

    divider(doc)

    cover_info = doc.add_paragraph()
    cover_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_info.paragraph_format.space_before = Pt(20)
    r = cover_info.add_run(
        f'Target Router : ZTE F670L (IndiHome)\n'
        f'Tanggal Dibuat : {datetime.date.today().strftime("%d %B %Y")}\n'
        f'Kelompok : 6 — Mata Kuliah Jaringan Komputer'
    )
    r.font.name = 'Arial'; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(100,116,139)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # DAFTAR ISI
    # ═══════════════════════════════════════════
    heading(doc, 'DAFTAR ISI', level=1, before=0)
    toc_items = [
        ('BAB I',   'Ringkasan Eksekutif & Gambaran Aplikasi', '3'),
        ('BAB II',  'Tujuan & Manfaat Aplikasi', '4'),
        ('BAB III', 'Technology Stack yang Digunakan', '5'),
        ('BAB IV',  'Arsitektur Sistem', '6'),
        ('BAB V',   'Cara Menjalankan Aplikasi', '8'),
        ('BAB VI',  'Mekanisme Pengumpulan Data', '9'),
        ('BAB VII', 'Algoritma Deteksi Anomali Z-Score', '12'),
        ('BAB VIII','Widget & Fitur Dashboard (Cara Kerja Tiap Komponen)', '15'),
        ('BAB IX',  'API Endpoints Backend', '21'),
        ('BAB X',   'Monitoring: Apa Saja yang Dipantau', '23'),
        ('BAB XI',  'Skenario Masalah & Penanganannya', '25'),
        ('BAB XII', 'Kelebihan & Kekurangan Sistem', '29'),
        ('BAB XIII','Kesimpulan', '31'),
    ]
    for bab, title, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        add_run(p, f'{bab}  ', bold=True, color=(37,99,235), size=10)
        add_run(p, f'{title}', size=10)
        add_run(p, f'  ...... {page}', italic=True, color=(148,163,184), size=9)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # BAB I — RINGKASAN EKSEKUTIF
    # ═══════════════════════════════════════════
    heading(doc, 'BAB I — RINGKASAN EKSEKUTIF & GAMBARAN APLIKASI', level=1)

    body(doc,
        'NetVision Pro adalah aplikasi monitoring jaringan nirkabel (WiFi) berbasis web yang dirancang '
        'untuk memantau kinerja router rumahan ZTE F670L secara real-time. Aplikasi ini dikembangkan '
        'menggunakan arsitektur full-stack modern dengan komponen backend berbasis Node.js, frontend '
        'berbasis React.js, dan database cloud menggunakan Supabase (PostgreSQL).'
    )
    body(doc,
        'Berbeda dengan aplikasi monitoring konvensional yang mengandalkan protokol standar seperti SNMP '
        '(Simple Network Management Protocol), NetVision Pro menggunakan teknik web scraping headless '
        'melalui library Puppeteer. Ini dikarenakan router ZTE F670L yang disediakan oleh IndiHome '
        'mengunci akses SNMP dan tidak menyediakan API publik, sehingga satu-satunya cara untuk '
        'membaca data bandwidth adalah dengan mengakses halaman admin web router dan mengekstrak '
        'data secara otomatis dari elemen HTML.'
    )
    body(doc,
        'Nilai tambah utama aplikasi ini adalah mesin deteksi anomali berbasis algoritma Z-Score '
        'Multi-Dimensi yang bersifat self-learning. Setiap 5 menit, model melatih ulang dirinya '
        'sendiri menggunakan 100 data historis terakhir dari database, sehingga batas pendeteksian '
        'anomali selalu adaptif terhadap pola penggunaan jaringan yang berubah setiap hari.'
    )

    heading(doc, '1.1 Identitas Aplikasi', level=2)
    two_col_table(doc, [
        ('Nama Aplikasi',     'NetVision Pro'),
        ('Versi',             '1.0.0'),
        ('Tipe Aplikasi',     'Web Dashboard — Network Monitoring & Anomaly Detection'),
        ('Target Hardware',   'Router ZTE F670L (IndiHome / ISP Indonesia)'),
        ('Mode Operasi',      'LIVE (scraping router nyata) atau MOCK (simulasi tanpa router)'),
        ('URL Publik',        'https://jarkom.clyuti.my.id (via Cloudflare Tunnel)'),
        ('Polling Interval',  'Setiap 30–60 detik (scraping router) + setiap 5 detik (frontend refresh)'),
        ('Database',          'Supabase Cloud PostgreSQL — Tabel: network_logs'),
        ('Kelompok',          'Kelompok 6 — Jaringan Komputer'),
    ])
    doc.add_paragraph()

    body(doc,
        'Aplikasi ini berjalan sepenuhnya di komputer lokal (localhost:5000) dan diekspos ke '
        'internet melalui Cloudflare Tunnel tanpa perlu konfigurasi port forwarding pada router ISP. '
        'Data yang ditampilkan di dashboard di-refresh setiap 5 detik oleh frontend, sementara '
        'proses pengambilan data dari router berjalan setiap 30 detik di latar belakang.'
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # BAB II — TUJUAN & MANFAAT
    # ═══════════════════════════════════════════
    heading(doc, 'BAB II — TUJUAN & MANFAAT APLIKASI', level=1)
    heading(doc, '2.1 Tujuan Pengembangan', level=2)
    tujuan = [
        ('Memantau Traffic Real-Time: ',
         'Memberikan visibilitas langsung terhadap kecepatan unduh (Rx/Download) dan unggah '
         '(Tx/Upload) pada jaringan WiFi secara terus-menerus.'),
        ('Mendeteksi Anomali Otomatis: ',
         'Menggunakan algoritma Z-Score multi-dimensi untuk mendeteksi lonjakan traffic yang '
         'tidak wajar (bandwidth spike) atau beban CPU berlebih tanpa intervensi manual.'),
        ('Melacak Klien WiFi: ',
         'Mengetahui siapa saja perangkat yang sedang terhubung ke router (hostname, IP, MAC '
         'Address, SSID, dan sejak kapan perangkat tersebut terhubung).'),
        ('Menyimpan Riwayat Jaringan: ',
         'Mencatat setiap snapshot kondisi jaringan ke database cloud Supabase sehingga '
         'administrator dapat menganalisis tren jangka panjang.'),
        ('Memberikan Skor Kesehatan Jaringan: ',
         'Mengkalkulasi Network Score (0–100) dan grade (A/B/C/D) secara real-time '
         'berdasarkan kondisi CPU dan status anomali yang aktif.'),
    ]
    for label, desc in tujuan:
        bullet(doc, label, desc)

    heading(doc, '2.2 Manfaat Praktis', level=2)
    manfaat = [
        ('Bagi Pengguna Rumah: ',
         'Dapat mendeteksi jika ada perangkat asing (tidak dikenal) yang masuk ke jaringan '
         'WiFi dan menggunakan bandwidth secara berlebihan.'),
        ('Bagi Administrator Jaringan: ',
         'Memiliki data historis yang terstruktur untuk analisis kapasitas jaringan dan '
         'perencanaan upgrade infrastruktur.'),
        ('Bagi Presentasi & Akademik: ',
         'Dashboard dapat diakses dari luar jaringan melalui URL publik, sangat berguna '
         'untuk demonstrasi di kampus tanpa koneksi ke router fisik.'),
        ('Bagi Pemeliharaan Proaktif: ',
         'Notifikasi visual anomali yang langsung terlihat di dashboard memungkinkan tindakan '
         'korektif sebelum masalah jaringan memburuk.'),
    ]
    for label, desc in manfaat:
        bullet(doc, label, desc)
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # BAB III — TECH STACK
    # ═══════════════════════════════════════════
    heading(doc, 'BAB III — TECHNOLOGY STACK YANG DIGUNAKAN', level=1)

    heading(doc, '3.1 Backend', level=2)
    multi_col_table(doc,
        ['Library / Tool', 'Versi', 'Fungsi dalam Sistem'],
        [
            ('Node.js',            '≥ 18 LTS',   'Runtime JavaScript sisi server — menjalankan seluruh logika backend'),
            ('Express.js',         '^4.19.2',    'Framework web server — menyediakan REST API dan melayani file statis frontend'),
            ('Puppeteer',          '^22.15.0',   'Headless Chromium browser — digunakan untuk login otomatis dan scraping halaman admin router ZTE F670L'),
            ('@supabase/supabase-js', '^2.43.4', 'Supabase JavaScript Client — menghubungkan backend ke database cloud Supabase PostgreSQL'),
            ('dotenv',             '^16.4.5',    'Memuat variabel konfigurasi dari file .env (ROUTER_PASS, SUPABASE_URL, dll.)'),
            ('cors',               '^2.8.5',     'Middleware CORS — mengizinkan frontend mengakses API backend dari origin berbeda'),
            ('child_process',      'built-in',   'Menjalankan perintah PowerShell untuk mengambil data CPU Load dan Uptime dari Windows'),
        ],
        col_widths=[2.0, 1.2, 3.5]
    )
    doc.add_paragraph()

    heading(doc, '3.2 Frontend', level=2)
    multi_col_table(doc,
        ['Library / Tool', 'Versi', 'Fungsi dalam Sistem'],
        [
            ('React.js',          '^19.0.0',  'Library UI — membangun antarmuka dashboard yang reaktif dan berbasis komponen'),
            ('Vite',              '^8.0.12',  'Build tool & dev server — kompilasi aset frontend ke folder dist/ untuk produksi'),
            ('Tailwind CSS',      '^3.4.4',   'Framework CSS utility-first — styling seluruh komponen dashboard (bento grid, glassmorphism, dll.)'),
            ('Recharts',          '^2.12.7',  'Library charting React — merender LineChart (Live Traffic & CPU History) dan PieChart (CPU Gauge)'),
            ('Axios',             '^1.7.2',   'HTTP client — melakukan fetch data dari API backend (/api/metrics/current, /api/metrics/history, /api/model/stats)'),
            ('Lucide React',      '^0.395.0', 'Library ikon SVG — menyediakan ikon-ikon visual (Activity, Wifi, Cpu, ArrowDown, Users, History, dll.)'),
        ],
        col_widths=[1.8, 1.2, 3.7]
    )
    doc.add_paragraph()

    heading(doc, '3.3 Database & Infrastruktur', level=2)
    multi_col_table(doc,
        ['Komponen', 'Teknologi', 'Peran'],
        [
            ('Cloud Database',  'Supabase (PostgreSQL)', 'Menyimpan seluruh log monitoring historis dalam tabel network_logs'),
            ('Tunneling',       'Cloudflare Tunnel (cloudflared)', 'Mengekspos localhost:5000 ke domain publik jarkom.clyuti.my.id tanpa port forwarding'),
            ('OS Layer',        'Windows 10/11 + PowerShell', 'Sumber data CPU Load (WMI Win32_Processor) dan System Uptime (Win32_OperatingSystem)'),
            ('Target Router',   'ZTE F670L (IndiHome)', 'Hardware sumber data bandwidth dan daftar klien WiFi yang di-scrape melalui halaman adminnya'),
        ],
        col_widths=[1.8, 2.0, 2.9]
    )
    doc.add_paragraph()
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # BAB IV — ARSITEKTUR SISTEM
    # ═══════════════════════════════════════════
    heading(doc, 'BAB IV — ARSITEKTUR SISTEM', level=1)
    body(doc,
        'Sistem NetVision Pro menggunakan arsitektur tiga lapis (three-tier architecture) yang terdiri '
        'dari lapisan presentasi (frontend), lapisan logika bisnis (backend), dan lapisan data (database). '
        'Keempat komponen utama berikut saling berinteraksi untuk menghasilkan monitoring jaringan '
        'yang real-time dan cerdas.'
    )

    heading(doc, '4.1 Alur Data Keseluruhan (Data Flow)', level=2)
    body(doc,
        'Berikut adalah urutan alur data dari sumber hingga tampil di layar pengguna:'
    )
    steps = [
        ('1 ▸ Polling Berkala (tiap 30 dtk): ',
         'Proses setInterval di backend (server.js) memanggil fungsi pollAndSave() yang mengotomatiskan '
         'seluruh siklus pengambilan dan penyimpanan data.'),
        ('2 ▸ Web Scraping Router: ',
         'Puppeteer headless Chrome (persistentPage singleton) membuka halaman admin router ZTE F670L '
         'di 192.168.1.1, login secara otomatis, navigasi ke menu WLAN Status, dan membaca nilai '
         'TotalBytesCount:N serta daftar klien dari elemen DOM.'),
        ('3 ▸ Pengambilan CPU (PowerShell): ',
         'Secara paralel, perintah PowerShell (execSync) dijalankan untuk membaca persentase beban '
         'CPU dari kelas WMI Win32_Processor dan durasi uptime dari Win32_OperatingSystem.'),
        ('4 ▸ Kalkulasi Kecepatan (Mbps): ',
         'Backend menghitung kecepatan jaringan dari selisih byte (Δbytes) saat ini vs siklus '
         'sebelumnya dibagi dengan selisih waktu (Δt): Speed = (Δbytes × 8) / (Δt × 10⁶) Mbps.'),
        ('5 ▸ Deteksi Anomali Z-Score: ',
         'Nilai CPU, Rx, dan Tx terbaru dibandingkan dengan model statistik (mean μ + stdDev σ) yang '
         'telah dilatih dari 100 log historis Supabase. Jika Z = |x−μ|/σ > 3.0 dan melebihi batas '
         'minimum fisik, data ditandai is_anomaly = true.'),
        ('6 ▸ Simpan ke Supabase: ',
         'Semua metrik (cpu_usage, rx_bytes, tx_bytes, is_anomaly, anomaly_reason, client_count, '
         'uptime) dikirim ke tabel network_logs di Supabase menggunakan Supabase JS Client.'),
        ('7 ▸ Perbarui State In-Memory: ',
         'Nilai-nilai terbaru disimpan ke variabel global currentState di memori backend agar '
         'dapat langsung diakses oleh API endpoint tanpa perlu query ulang ke database.'),
        ('8 ▸ Frontend Polling API (tiap 5 dtk): ',
         'React frontend memanggil /api/metrics/current (live state) dan /api/metrics/history '
         '(50 log terakhir dari Supabase) menggunakan axios setiap 5 detik.'),
        ('9 ▸ Render Dashboard: ',
         'Komponen Dashboard.jsx memperbarui state (useState) dengan data baru, memicu re-render '
         'otomatis seluruh widget — grafik, angka kecepatan, indikator status, daftar klien, dsb.'),
    ]
    for label, desc in steps:
        bullet(doc, label, desc)

    heading(doc, '4.2 Struktur Direktori Proyek', level=2)
    multi_col_table(doc,
        ['Path', 'Teknologi', 'Keterangan'],
        [
            ('backend/server.js',         'Node.js, Express, Puppeteer', 'Titik masuk backend — berisi seluruh logika polling, scraping, Z-Score, API endpoints, dan server start'),
            ('backend/.env',              'dotenv',                      'Konfigurasi sensitif: SUPABASE_URL, SUPABASE_KEY, ROUTER_IP, ROUTER_USER, ROUTER_PASS, PORT, MOCK_MODE, ANOMALY_THRESHOLD'),
            ('backend/package.json',      'npm',                         'Manifest dependensi backend'),
            ('frontend/src/components/Dashboard.jsx', 'React, Recharts','Komponen utama UI — berisi seluruh widget, logika fetch data, dan rendering dashboard'),
            ('frontend/src/index.css',    'Tailwind CSS',                'Styling global, tema light glassmorphism, konfigurasi font'),
            ('frontend/dist/',            'Vite build output',           'Folder aset produksi (HTML/JS/CSS terminifikasi) yang di-serve oleh Express backend'),
            ('frontend/vite.config.js',   'Vite',                        'Konfigurasi proxy API: request ke /api diteruskan ke localhost:5000'),
            ('frontend/package.json',     'npm',                         'Manifest dependensi frontend'),
        ],
        col_widths=[2.3, 1.8, 2.6]
    )
    doc.add_paragraph()
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # BAB V — CARA MENJALANKAN
    # ═══════════════════════════════════════════
    heading(doc, 'BAB V — CARA MENJALANKAN APLIKASI', level=1)

    heading(doc, '5.1 Prasyarat Sistem', level=2)
    multi_col_table(doc,
        ['Komponen', 'Versi Minimum', 'Keterangan'],
        [
            ('Node.js',          '18 LTS',     'Runtime JavaScript untuk menjalankan backend server'),
            ('npm',              '9.x',         'Package manager — terinstal bersama Node.js'),
            ('Google Chrome',    'Terinstall',  'Diperlukan oleh Puppeteer untuk menjalankan headless browser'),
            ('Python 3',         '3.8+',        'Diperlukan untuk menjalankan script generator dokumentasi (generate_full_doc.py)'),
            ('python-docx',      'pip install', 'Library Python untuk membuat file Word dari script'),
            ('Cloudflared',      'Terbaru',     'Biner Cloudflare Tunnel untuk mengekspos server ke internet publik'),
            ('Koneksi WiFi',     'Router ZTE F670L', 'Komputer yang menjalankan backend HARUS terhubung ke WiFi dengan router target'),
        ],
        col_widths=[1.8, 1.5, 3.4]
    )
    doc.add_paragraph()

    heading(doc, '5.2 Langkah Instalasi & Setup', level=2)

    steps_install = [
        ('Langkah 1 — Konfigurasi File .env',
         r'd:\jarkom\backend\.env',
         ['PORT=5000',
          'SUPABASE_URL=https://[project-id].supabase.co',
          'SUPABASE_KEY=[anon public key dari Supabase Dashboard]',
          'ROUTER_IP=192.168.1.1',
          'ROUTER_USER=admin',
          'ROUTER_PASS=[password router Anda]',
          'MOCK_MODE=false',
          'ANOMALY_THRESHOLD=3.0']),
        ('Langkah 2 — Install Dependensi Backend',
         r'd:\jarkom\backend',
         ['npm install']),
        ('Langkah 3 — Install Dependensi Frontend & Build',
         r'd:\jarkom\frontend',
         ['npm install', 'npm run build']),
        ('Langkah 4 — Jalankan Backend Server',
         r'd:\jarkom\backend',
         ['npm start',
          '# Output sukses:',
          '# ✅ Supabase client initialized.',
          '# 🚀 Backend berjalan di http://localhost:5000',
          '# [TRAINING] ✅ Model dilatih dari 100 record.']),
        ('Langkah 5 — Jalankan Cloudflare Tunnel (opsional)',
         'Semua direktori',
         [r'& "C:\Program Files (x86)\cloudflared\cloudflared.exe" tunnel run jarkom']),
    ]

    for step_title, cwd, cmds in steps_install:
        heading(doc, step_title, level=3)
        if cwd and not cwd.startswith('#'):
            p = doc.add_paragraph()
            set_para_spacing(p, 0, 4)
            add_run(p, f'  📁  Direktori: ', bold=True, color=(71,85,105), size=10)
            add_run(p, cwd, color=(37,99,235), size=10, font='Consolas')
        code_block(doc, cmds)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    heading(doc, '5.3 Mengakses Dashboard', level=2)
    two_col_table(doc, [
        ('Akses Lokal',   'Buka browser → http://localhost:5000'),
        ('Akses Publik',  'Buka browser → https://jarkom.clyuti.my.id (memerlukan Cloudflare Tunnel aktif)'),
        ('Health Check',  'GET http://localhost:5000/api/health → mengembalikan JSON status server'),
        ('Mode MOCK',     'Set MOCK_MODE=true di .env untuk menjalankan simulasi tanpa koneksi router fisik'),
    ])
    doc.add_paragraph()
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # BAB VI — MEKANISME PENGUMPULAN DATA
    # ═══════════════════════════════════════════
    heading(doc, 'BAB VI — MEKANISME PENGUMPULAN DATA', level=1)
    body(doc,
        'NetVision Pro mengumpulkan data dari dua sumber berbeda: (1) halaman admin web router ZTE F670L '
        'melalui Puppeteer headless browser untuk data bandwidth dan klien WiFi, dan (2) perintah '
        'PowerShell Windows untuk data CPU dan uptime sistem.'
    )

    heading(doc, '6.1 Web Scraping Router ZTE F670L (Puppeteer)', level=2)
    body(doc,
        'Router ZTE F670L yang digunakan ISP IndiHome mengunci akses SNMP dan tidak menyediakan API '
        'terbuka. Satu-satunya cara untuk membaca data bandwidth adalah melalui halaman admin web '
        'pada alamat http://192.168.1.1. NetVision Pro mengotomatiskan proses ini menggunakan '
        'Puppeteer — library yang mengendalikan browser Chrome secara programatik (headless, tanpa '
        'antarmuka visual).'
    )

    heading(doc, '6.1.1 Page Singleton Pattern (Optimasi Performa)', level=3)
    body(doc,
        'Implementasi krusial yang membedakan NetVision Pro dari solusi scraping sederhana adalah '
        'penggunaan "persistent page singleton". Alih-alih membuat tab browser baru di setiap '
        'siklus polling (yang membutuhkan 15–20 detik per siklus), sistem mempertahankan satu tab '
        'browser yang tetap hidup dan hanya melakukan navigasi ulang yang minimal.'
    )
    bullet(doc, 'browserInstance: ',
           'Satu instance browser Chrome headless yang dibuat saat startup dan digunakan kembali '
           'di setiap siklus. Hanya dibuat ulang jika crash terdeteksi.')
    bullet(doc, 'persistentPage: ',
           'Satu tab halaman yang tetap aktif. Sistem memeriksa apakah tab masih hidup dengan '
           'memanggil page.title(). Jika tab crash, dibuat ulang.')
    bullet(doc, 'isPageReady: ',
           'Flag boolean yang menandakan apakah tab saat ini sudah berada di halaman WLAN Status '
           'dan siap untuk scraping tanpa navigasi penuh.')
    body(doc, 'Keuntungan: Waktu scraping turun dari 15–20 detik menjadi ~3–5 detik per siklus.', italic=True)

    heading(doc, '6.1.2 Alur Scraping Step-by-Step', level=3)
    scrape_steps = [
        ('Step 1 — Buka Halaman Utama Router: ',
         'Puppeteer membuka http://192.168.1.1/ dan menunggu DOM selesai dimuat '
         '(waitUntil: "domcontentloaded").'),
        ('Step 2 — Deteksi & Eksekusi Login: ',
         'Sistem memeriksa apakah elemen input[type="password"] ada. Jika ada, berarti sesi '
         'sudah habis dan login perlu dilakukan ulang. Kredensial diisi ke elemen #Frm_Username '
         'dan #Frm_Password, lalu tombol #LoginId diklik.'),
        ('Step 3 — Navigasi Menu: ',
         'Puppeteer mengklik menu "Local Network" (#localnet), lalu "Status" (#localNetStatus), '
         'dan collapsible bar "WLAN Status" (#WLANStatusBar). Setiap klik diberi jeda 1.5 detik '
         'untuk memungkinkan AJAX router selesai memuat data.'),
        ('Step 4 — Klik Tombol Refresh Router: ',
         'Tombol refresh (#WLANStatus_Btn_refresh) diklik agar router memperbarui counter '
         'TotalBytesCount sebelum dibaca. Tanpa ini, nilai bytes tidak akan berubah antar siklus.'),
        ('Step 5 — Ekspansi Panel Klien: ',
         'Bar "WLAN Client Status" (#Wlan_ClientStatBar) diklik untuk memuat daftar perangkat '
         'yang terhubung, dilanjutkan klik tombol refresh klien (#Btn_refresh_Wlan_ClientStat).'),
        ('Step 6 — Ekstraksi Data Bytes: ',
         'Menggunakan page.evaluate() untuk menjalankan kode JavaScript langsung di dalam browser. '
         'Sistem mengiterasi elemen TotalBytesCount:0 hingga TotalBytesCount:7 (satu per SSID: '
         '2.4GHz, 5GHz, dsb.), membaca atribut title yang berformat "rxBytes/txBytes", '
         'mengakumulasikan semua SSID untuk mendapatkan total bytes.'),
        ('Step 7 — Ekstraksi Daftar Klien: ',
         'Mengiterasi semua elemen [id^="template_Wlan_ClientStat_"] yang tidak tersembunyi. '
         'Untuk setiap klien, membaca HostName:N, IPAddress:N, MACAddress:N, dan ESSID:N.'),
        ('Step 8 — Kalkulasi Kecepatan: ',
         'Backend menghitung delta bytes (selisih dengan siklus sebelumnya) dan membaginya '
         'dengan selisih waktu untuk mendapatkan kecepatan dalam Mbps. Guard tambahan '
         'memastikan nilai negatif (akibat router reboot dan counter reset) tidak ditampilkan.'),
    ]
    for label, desc in scrape_steps:
        bullet(doc, label, desc)

    heading(doc, '6.2 Pengambilan Data CPU & Uptime (PowerShell)', level=2)
    body(doc,
        'Data performa sistem operasi diambil menggunakan modul child_process bawaan Node.js '
        '(execSync), yang menjalankan perintah PowerShell secara sinkron dengan timeout 3 detik.'
    )
    body(doc, 'Perintah PowerShell untuk CPU Load:')
    code_block(doc, [
        "Get-CimInstance Win32_Processor | Measure-Object -Property LoadPercentage -Average",
        "| Select-Object -ExpandProperty Average"
    ])
    body(doc, 'Perintah PowerShell untuk System Uptime:')
    code_block(doc, [
        "(Get-Date) - (gcim Win32_OperatingSystem).LastBootUpTime",
        "| Select-Object -ExpandProperty TotalSeconds"
    ])
    body(doc,
        'Jika PowerShell gagal (timeout atau error), sistem menggunakan fallback: '
        'process.uptime() untuk uptime Node.js dan nilai acak 10–25% untuk CPU sebagai '
        'indikator tidak tersedia.',
        italic=True
    )

    heading(doc, '6.3 Database Supabase (network_logs)', level=2)
    body(doc,
        'Setiap siklus polling yang berhasil, backend menyimpan satu baris record ke tabel '
        'network_logs di Supabase Cloud PostgreSQL.'
    )
    multi_col_table(doc,
        ['Nama Kolom', 'Tipe Data', 'Keterangan'],
        [
            ('id',             'uuid (auto)',  'Primary key, dibuat otomatis oleh Supabase'),
            ('created_at',     'timestamptz', 'Waktu pencatatan, dibuat otomatis oleh Supabase'),
            ('router_ip',      'text',         'IP address router yang dimonitor (192.168.1.1)'),
            ('cpu_usage',      'int4',         'Persentase beban CPU server (0–100)'),
            ('uptime',         'text',         'Durasi uptime sistem dalam format "Nh Nm Ns"'),
            ('rx_bytes',       'text',         'Total byte kumulatif diterima (disimpan sebagai string karena BigInt)'),
            ('tx_bytes',       'text',         'Total byte kumulatif dikirim (disimpan sebagai string karena BigInt)'),
            ('client_count',   'int4',         'Jumlah klien WiFi yang terdeteksi pada siklus ini'),
            ('is_anomaly',     'boolean',      'True jika Z-Score melebihi ambang batas pada siklus ini'),
            ('anomaly_reason', 'text',         'Alasan anomali dalam kalimat (null jika tidak anomali)'),
        ],
        col_widths=[1.6, 1.3, 3.8]
    )
    doc.add_paragraph()
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # BAB VII — ALGORITMA Z-SCORE
    # ═══════════════════════════════════════════
    heading(doc, 'BAB VII — ALGORITMA DETEKSI ANOMALI Z-SCORE', level=1)
    body(doc,
        'Inti kecerdasan NetVision Pro terletak pada implementasi algoritma Z-Score Multi-Dimensi '
        'yang bersifat dinamis (self-learning). Algoritma ini secara otomatis menyesuaikan '
        'batas pendeteksian anomali berdasarkan perilaku historis jaringan, sehingga mengurangi '
        'false alarm yang umum terjadi pada sistem dengan threshold statis.'
    )

    heading(doc, '7.1 Definisi Z-Score', level=2)
    body(doc,
        'Z-Score mengukur seberapa jauh suatu nilai menyimpang dari rata-rata populasinya, '
        'dinyatakan dalam satuan standar deviasi. Rumus matematisnya:'
    )
    p_formula = doc.add_paragraph()
    p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_formula.paragraph_format.space_before = Pt(8)
    p_formula.paragraph_format.space_after  = Pt(8)
    r_f = p_formula.add_run('Z = | X − μ | / σ')
    r_f.font.name = 'Consolas'; r_f.font.size = Pt(16); r_f.bold = True
    r_f.font.color.rgb = RGBColor(37,99,235)

    two_col_table(doc, [
        ('X (X)',    'Nilai metrik real-time yang sedang dievaluasi (CPU%, Rx Mbps, atau Tx Mbps)'),
        ('μ (Mu)',   'Nilai rata-rata (mean) dari 100 log historis terakhir di Supabase'),
        ('σ (Sigma)','Standar deviasi dari 100 log historis terakhir di Supabase'),
        ('Threshold','ANOMALY_THRESHOLD = 3.0 (dapat dikonfigurasi melalui .env)'),
    ])
    doc.add_paragraph()

    heading(doc, '7.2 Proses Training Dinamis (Self-Learning)', level=2)
    body(doc,
        'Model statistik dilatih ulang secara otomatis setiap 5 menit (setInterval setiap 300.000 ms). '
        'Proses training mengambil 100 log terbaru dari Supabase dan menghitung statistik:'
    )
    bullet(doc, 'Mean (μ): ', 'Rata-rata dari semua nilai CPU, Rx speed, dan Tx speed dalam 100 log tersebut.')
    bullet(doc, 'Standard Deviation (σ): ',
           'Menggunakan rumus deviasi standar populasi: σ = √(Σ(xᵢ−μ)² / N). '
           'Jika σ = 0 (semua data identik), nilai 1 digunakan sebagai fallback untuk mencegah pembagian dengan nol.')
    body(doc,
        'Dengan re-training setiap 5 menit, model beradaptasi terhadap pola harian: '
        'traffic rendah di malam hari, tinggi di siang hari. Ini mencegah alarm palsu '
        'di jam-jam tertentu yang memiliki pola traffic normal yang berbeda.',
        italic=True
    )

    heading(doc, '7.3 Kondisi Anomali (Dual-Guard System)', level=2)
    body(doc,
        'Untuk mencegah false alarm pada kondisi traffic sangat rendah (di mana δ kecil apapun '
        'bisa menghasilkan Z-Score besar), sistem menerapkan dual-guard condition: sebuah metrik '
        'hanya dianggap anomali jika KEDUANYA dipenuhi:'
    )
    multi_col_table(doc,
        ['Metrik', 'Guard 1: Z-Score', 'Guard 2: Nilai Minimum Fisik', 'Contoh Anomali Dipicu'],
        [
            ('CPU Load',        'Z > 3.0', 'CPU > 60%',       'CPU = 95% saat μ=46%, σ=15%  →  Z=3.27 ✓'),
            ('Download (Rx)',   'Z > 3.0', 'Rx > 50 Mbps',    'Rx = 200Mbps saat μ=2Mbps, σ=26Mbps  →  Z=7.6 ✓'),
            ('Upload (Tx)',     'Z > 3.0', 'Tx > 20 Mbps',    'Tx = 80Mbps saat μ=3Mbps, σ=15Mbps  →  Z=5.1 ✓'),
        ],
        col_widths=[1.5, 1.3, 2.0, 2.5]
    )
    doc.add_paragraph()

    heading(doc, '7.4 Cold Start Fallback (Perlindungan Awal)', level=2)
    body(doc,
        'Ketika database memiliki kurang dari 20 record (kondisi awal sistem baru dipasang), '
        'standar deviasi belum stabil untuk perhitungan statistik yang akurat. Sistem secara '
        'otomatis beralih ke mode threshold statis konservatif:'
    )
    multi_col_table(doc,
        ['Metrik', 'Threshold Statis (Cold Start)'],
        [
            ('CPU Load',      'CPU > 85%'),
            ('Download (Rx)', 'Rx > 150 Mbps'),
            ('Upload (Tx)',   'Tx > 50 Mbps'),
        ],
        col_widths=[2.0, 4.0]
    )
    doc.add_paragraph()

    heading(doc, '7.5 Output Deteksi Anomali', level=2)
    body(doc,
        'Hasil deteksi anomali disimpan ke dua field di currentState (in-memory) dan ke database:'
    )
    code_block(doc, [
        '// Contoh output saat anomali terdeteksi:',
        '{',
        '  is_anomaly: true,',
        '  anomaly_reason: "CPU Load abnormal: 95% (normal ~46%) & Download abnormal: 180Mbps (normal ~2Mbps)"',
        '}',
        '',
        '// Contoh output kondisi normal:',
        '{',
        '  is_anomaly: false,',
        '  anomaly_reason: null',
        '}',
    ])
    doc.add_paragraph()
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # BAB VIII — WIDGET & FITUR DASHBOARD
    # ═══════════════════════════════════════════
    heading(doc, 'BAB VIII — WIDGET & FITUR DASHBOARD (CARA KERJA TIAP KOMPONEN)', level=1)
    body(doc,
        'Antarmuka NetVision Pro dibangun menggunakan konsep Bento Grid — tata letak kartu-kartu '
        'informasi berukuran berbeda yang tersusun rapi, terinspirasi dari desain modern Apple dan '
        'Linear. Setiap kartu menggunakan efek glassmorphism (transparan dengan blur) pada tema '
        'light mode dengan aksen warna biru-indigo.'
    )
    body(doc,
        'Frontend React.js melakukan polling data dari backend setiap 5 detik menggunakan '
        'setInterval. Setiap respons baru langsung memperbarui state React (useState/useMemo), '
        'yang secara otomatis memicu re-render seluruh widget yang bergantung pada data tersebut.'
    )

    widgets = [
        ('Widget 1: Header Bar (Status Sistem)',
         'Terletak di bagian paling atas dashboard. Berisi: (a) nama aplikasi "NetVision Pro" '
         'dengan ikon gradien biru-indigo; (b) target IP router yang dipantau; (c) indikator '
         'status koneksi dengan animasi ping — hijau+berdenyut (LIVE), ungu (MOCK MODE), atau '
         'merah (OFFLINE); (d) badge jumlah klien WiFi aktif (X Client); dan (e) tombol manual '
         'refresh (ikon putar) yang memicu fetchData() secara langsung tanpa menunggu interval 5 detik.',
         'Sumber data: currentMetrics.is_online, is_mock, client_count, router_ip. '
         'Status "LIVE" hanya muncul jika scraping router berhasil dan is_mock = false.',
         'Jika backend offline → status OFFLINE (merah). Jika MOCK_MODE=true → status MOCK MODE (ungu).'),

        ('Widget 2: Download Speed (Rx)',
         'Menampilkan kecepatan unduh real-time saat ini dalam satuan Mbps dengan angka besar '
         'berwarna cyan. Di bawahnya terdapat progress bar horizontal yang bergerak proporsional '
         'dari 0 hingga 100 Mbps (cap di 100 Mbps). Di bagian bawah kartu ditampilkan total byte '
         'kumulatif yang telah diterima router sejak menyala, dikonversi ke Gigabyte (GB).',
         'Sumber data: currentMetrics.rx_speed_mbps (dikalkulasi dari Δbytes/Δt oleh backend), '
         'currentMetrics.rx_bytes (akumulasi total dari router sejak boot).',
         'Jika traffic 0 Mbps dalam waktu lama → kemungkinan router tidak mengirim data atau '
         'TotalBytesCount tidak berubah (counter stagnant issue).'),

        ('Widget 3: Upload Speed (Tx)',
         'Analog dengan widget Download, namun menampilkan kecepatan unggah dengan warna fuchsia '
         '(magenta). Progress bar menggunakan skala 0–50 Mbps. Total TX dalam GB ditampilkan '
         'di bagian bawah kartu.',
         'Sumber data: currentMetrics.tx_speed_mbps dan currentMetrics.tx_bytes.',
         'Nilai 0 Mbps terus-menerus dapat menandakan router hanya menerima traffic (skenario '
         'streaming/download murni) atau ada masalah scraping.'),

        ('Widget 4: System Uptime',
         'Menampilkan berapa lama sistem operasi komputer server (Windows) telah berjalan tanpa '
         'restart, dalam format "Jam h Menit m Detik s" (contoh: 174h 30m 19s). Widget ini '
         'bermanfaat untuk memantau stabilitas server. Di bagian bawah kartu ditampilkan '
         'keterangan interval polling yang aktif.',
         'Sumber data: currentMetrics.uptime — dikalkulasi dari PowerShell Win32_OperatingSystem.LastBootUpTime.',
         'Jika uptime tiba-tiba turun drastis → komputer server mengalami restart tidak terduga. '
         'Jika format "0h 0m 0s" → PowerShell gagal mengambil data uptime.'),

        ('Widget 5: CPU Load Gauge Chart',
         'Menampilkan persentase beban prosesor server dalam bentuk PieChart setengah lingkaran '
         '(semi-donut). Area biru menunjukkan CPU yang terpakai, area abu-abu transparan '
         'menunjukkan CPU yang tersedia. Di tengah grafik, persentase CPU ditampilkan secara '
         'numerik dalam teks tebal. Di bawah chart terdapat badge status tiga level: '
         '"Normal" (hijau, ≤50%), "Beban Sedang" (amber, >50%), dan "Kritis" (merah, >85%). '
         'Widget ini memiliki row-span 2 (dua baris tinggi) untuk memberikan ruang visual lebih.',
         'Sumber data: currentMetrics.cpu_usage dari PowerShell Get-CimInstance Win32_Processor. '
         'Diproses oleh useMemo cpuPieData agar tidak dihitung ulang setiap render.',
         'CPU spike sementara normal saat Windows Update atau antivirus berjalan. '
         'CPU tinggi konsisten dapat mengindikasikan beban Puppeteer berlebih atau proses lain.'),

        ('Widget 6: Live Traffic Chart (Bandwidth)',
         'Grafik LineChart besar (col-span 3) yang memvisualisasikan tren kecepatan jaringan '
         'secara historis. Dua garis paralel — biru (Download/Rx) dan fuchsia (Upload/Tx) — '
         'bergerak dari kiri ke kanan mengikuti waktu. Pada titik data yang terdeteksi sebagai '
         'anomali, CustomDot merender lingkaran merah berkedip (animate-ping) sebagai penanda '
         'visual peringatan. Tooltip interaktif muncul saat kursor diarahkan ke grafik, '
         'menampilkan detail waktu dan nilai Mbps. Sumbu X menampilkan waktu (HH:MM:SS) '
         'dan sumbu Y menampilkan nilai dalam Mbps.',
         'Sumber data: historyData — array 50 record dari /api/metrics/history (Supabase). '
         'Kecepatan Rx/Tx di-kalkulasi ulang oleh backend API dari selisih rx_bytes antar record.',
         'Grafik kosong di awal sistem → database belum memiliki record historis. '
         'Semua nilai 0 Mbps → masalah scraping atau TotalBytesCount tidak berubah.'),

        ('Widget 7: Network Score Card',
         'Kartu dinamis yang mengkalkulasi skor kesehatan jaringan real-time dari 0 hingga 100 '
         'poin dan mengkonversinya ke grade huruf (A/B/C/D). Algoritma scoring: mulai dari 100, '
         'kurangi 5 poin jika CPU>50%, kurangi 15 poin jika CPU>70%, kurangi 30 poin jika '
         'CPU>90%, dan kurangi 40 poin jika ada anomali aktif. Warna kartu berubah secara '
         'dinamis: hijau (Grade A, skor ≥90), biru (Grade B, ≥75), amber (Grade C, ≥50), '
         'merah dengan efek berkedip (Grade D, <50). Di bagian bawah kartu ditampilkan '
         'alasan anomali jika ada, atau "Sistem Berjalan Optimal" jika normal.',
         'Sumber data: currentMetrics.cpu_usage dan currentMetrics.is_anomaly. '
         'Dikalkulasi di sisi frontend oleh fungsi getNetworkScore() setiap re-render.',
         'Kartu berkedip merah terus-menerus saat Grade D → anomali aktif yang persisten '
         'perlu diselidiki segera (bandwidth spike atau CPU overload).'),

        ('Widget 8: CPU Load History Chart',
         'Grafik LineChart terpisah di baris kedua dashboard yang khusus memvisualisasikan '
         'tren beban CPU secara historis. Sumbu Y dibatasi 0–100%, sumbu X menampilkan '
         'timestamp. Garis biru (Blue-600) menunjukkan naik-turunnya CPU Load seiring waktu. '
         'Grafik ini membantu mengidentifikasi pola — apakah CPU spike terjadi bersamaan '
         'dengan lonjakan traffic atau secara independen.',
         'Sumber data: historyData[].cpu_usage — dari kolom cpu_usage tabel network_logs Supabase.',
         'Grafik datar di 0% → data CPU tidak terekam ke Supabase atau PowerShell gagal '
         'mengambil data sejak awal server berjalan.'),

        ('Widget 9: Metode Deteksi (Model Stats)',
         'Panel informasi yang menjelaskan metode Z-Score yang digunakan. Jika model sudah '
         'dilatih (modelStats.is_trained = true), panel menampilkan nilai μ/σ yang digunakan '
         'model saat ini untuk CPU, Rx, dan Tx, beserta Z-Score aktif saat ini untuk ketiga '
         'metrik. Nilai Z-Score yang melebihi threshold 3.0 ditampilkan dalam teks merah tebal '
         'sebagai indikator visual. Di bawahnya selalu ditampilkan informasi threshold yang '
         'aktif. Panel ini sangat berguna saat demonstrasi akademis untuk membuktikan bahwa '
         'sistem benar-benar menggunakan perhitungan statistik real-time.',
         'Sumber data: modelStats dari endpoint /api/model/stats — mengekspos cpuStats, '
         'rxStats, txStats, dan current_z_scores yang dikalkulasi dari currentState terkini.',
         'Jika panel menampilkan "Model dilatih setiap 5 menit..." → model belum selesai '
         'training pertama kali atau Supabase tidak terhubung.'),

        ('Widget 10: Client Terkoneksi (WLAN Clients)',
         'Panel daftar klien WiFi yang saat ini terhubung ke router. Setiap klien ditampilkan '
         'sebagai kartu dengan: avatar lingkaran biru (ikon WiFi), nama hostname perangkat '
         '(contoh: ASUS-ROG-Laptop), badge SSID (Winaila 5G / 2.4G), alamat IP lokal, '
         'MAC Address dalam format monospace, dan informasi "Since HH:MM" yang menunjukkan '
         'kapan perangkat tersebut pertama kali terdeteksi dalam sesi server ini. '
         'Panel dapat di-scroll jika jumlah klien banyak (max-height 280px).',
         'Sumber data: currentMetrics.clients — array dari scraping elemen '
         'template_Wlan_ClientStat_N router. active_since dikelola oleh Map clientFirstSeen '
         'di memori backend (menggunakan MAC address sebagai kunci).',
         'Jika "Tidak ada client WiFi terdeteksi" — panel Wlan_ClientStatBar router gagal '
         'diexpand atau tidak ada perangkat terhubung. active_since hanya tersedia sejak '
         'server terakhir kali di-restart (tidak persisten ke database).'),

        ('Widget 11: Log Database Historis (Supabase)',
         'Menampilkan 12 log terbaru dari Supabase sebagai daftar kartu berurutan. Setiap '
         'kartu log menampilkan: timestamp (HH:MM:SS + tanggal), badge status NORMAL (hijau) '
         'atau ANOMALI (merah), nilai CPU%, kecepatan Rx (↓ Mbps), dan kecepatan Tx (↑ Mbps). '
         'Kartu log anomali mendapatkan latar belakang merah muda pucat (bg-red-50) untuk '
         'membedakannya secara visual. Panel dapat di-scroll untuk melihat log lebih lama.',
         'Sumber data: historyData dari /api/metrics/history — 50 record terakhir dari Supabase, '
         'ditampilkan 12 terbaru (dibalik dari asc ke desc untuk urutan terbaru di atas).',
         'Semua Rx/Tx menampilkan 0.0 → kecepatan dikalkulasi dari delta bytes. Jika record '
         'pertama selalu 0 karena tidak ada data sebelumnya untuk dibandingkan, ini normal.'),
    ]

    for w_title, w_desc, w_data, w_issues in widgets:
        heading(doc, w_title, level=2)
        body(doc, '📋 Deskripsi Fungsi:', before=0, after=4)
        body(doc, w_desc, before=0, after=6)
        body(doc, '📊 Sumber & Alur Data:', before=0, after=4)
        body(doc, w_data, before=0, after=6, italic=True)
        body(doc, '⚠️  Potensi Masalah & Indikator:', before=0, after=4)
        body(doc, w_issues, before=0, after=10, italic=True)
        divider(doc)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # BAB IX — API ENDPOINTS
    # ═══════════════════════════════════════════
    heading(doc, 'BAB IX — API ENDPOINTS BACKEND', level=1)
    body(doc,
        'Backend Express.js menyediakan empat REST API endpoint yang dapat diakses oleh '
        'frontend maupun pihak ketiga. Seluruh respons dalam format JSON.'
    )
    multi_col_table(doc,
        ['Endpoint', 'Method', 'Keterangan', 'Respons Utama'],
        [
            ('/api/metrics/current', 'GET',
             'Mengembalikan state jaringan terkini dari memori backend (tanpa query DB)',
             'cpu_usage, rx_speed_mbps, tx_speed_mbps, uptime, clients[], is_anomaly, anomaly_reason, client_count, is_online, timestamp'),
            ('/api/metrics/history', 'GET',
             'Mengambil 50 log terakhir dari Supabase, menghitung kecepatan Rx/Tx dari delta bytes antar record',
             'Array record dengan rx_speed_mbps, tx_speed_mbps, cpu_usage, is_anomaly, created_at'),
            ('/api/model/stats',     'GET',
             'Mengekspos statistik model Z-Score: mean & stdDev tiap metrik, Z-Score saat ini, status training',
             'is_trained, threshold, cpu{mean,stdDev}, rx{...}, tx{...}, current_z_scores{cpu,rx,tx}'),
            ('/api/health',          'GET',
             'Health check endpoint — berguna untuk memverifikasi Cloudflare Tunnel aktif dan server merespons',
             'status:"ok", mode, router_ip, uptime, timestamp'),
        ],
        col_widths=[1.8, 0.9, 2.5, 2.5]
    )
    doc.add_paragraph()
    body(doc, 'Selain API endpoint, backend juga melayani file statis frontend:')
    code_block(doc, [
        '// server.js — Serve frontend static files',
        "app.use(express.static(path.join(__dirname, '../frontend/dist')));",
        "app.get('*', (req, res) => {",
        "  res.sendFile(path.join(__dirname, '../frontend/dist/index.html'));",
        "});",
    ])
    body(doc,
        'Ini berarti satu proses Node.js (npm start di /backend) sudah melayani '
        'BOTH API backend DAN file HTML/CSS/JS frontend. Tidak diperlukan server terpisah '
        'untuk frontend di mode produksi.',
        italic=True
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # BAB X — MONITORING: APA SAJA YANG DIPANTAU
    # ═══════════════════════════════════════════
    heading(doc, 'BAB X — MONITORING: APA SAJA YANG DIPANTAU', level=1)
    body(doc,
        'NetVision Pro memantau enam kategori metrik yang saling melengkapi untuk memberikan '
        'gambaran lengkap kondisi jaringan dan server:'
    )
    multi_col_table(doc,
        ['Kategori', 'Metrik', 'Satuan', 'Sumber', 'Interval'],
        [
            ('Bandwidth Download',   'Rx Speed (kecepatan unduh real-time)', 'Mbps', 'Router ZTE — TotalBytesCount delta', '30 dtk'),
            ('Bandwidth Upload',     'Tx Speed (kecepatan unggah real-time)', 'Mbps', 'Router ZTE — TotalBytesCount delta', '30 dtk'),
            ('Volume Kumulatif',     'Total Rx Bytes & Tx Bytes sejak router nyala', 'GB', 'Router ZTE — TotalBytesCount raw', '30 dtk'),
            ('Beban Prosesor',       'CPU Load percentage', '%', 'PowerShell Win32_Processor', '30 dtk'),
            ('Stabilitas Server',    'System Uptime (durasi server nyala)', 'Jam/Menit/Detik', 'PowerShell Win32_OperatingSystem', '30 dtk'),
            ('Perangkat Jaringan',   'Jumlah & detail klien WiFi terhubung', 'Hostname, IP, MAC, SSID', 'Router ZTE — WLAN Client Status', '30 dtk'),
            ('Kesehatan Jaringan',   'Network Score dan Grade (A/B/C/D)', 'Skor 0–100', 'Kalkulasi frontend dari CPU & anomali', '5 dtk'),
            ('Status Anomali',       'Apakah ada lonjakan traffic / CPU abnormal', 'Boolean + alasan teks', 'Algoritma Z-Score backend', '30 dtk'),
            ('Model Statistik',      'Mean (μ) dan StdDev (σ) tiap metrik', 'Float', 'Training dari 100 log Supabase', '5 menit'),
            ('Z-Score Aktif',        'Jarak deviasi tiap metrik dari normalnya', 'Dimensionless', 'Kalkulasi backend dari model stats', '30 dtk'),
        ],
        col_widths=[1.8, 2.2, 1.3, 1.8, 0.8]
    )
    doc.add_paragraph()

    heading(doc, '10.1 Yang TIDAK Dipantau (Batasan Sistem)', level=2)
    body(doc, 'Berikut adalah metrik yang berada di luar cakupan monitoring NetVision Pro saat ini:')
    limitations = [
        'Kecepatan internet ke ISP (WAN speed test) — hanya traffic internal WiFi yang dimonitor',
        'Penggunaan bandwidth per-perangkat (per-client bandwidth) — router ZTE tidak menyediakan data ini melalui halaman admin',
        'Temperatur router — tidak tersedia di halaman admin ZTE F670L',
        'Latensi (ping/RTT) — tidak diukur dalam siklus scraping saat ini',
        'Paket loss — tidak tersedia melalui web scraping halaman admin router',
        'QoS (Quality of Service) settings — hanya monitoring, tidak ada fitur kontrol',
        'Traffic per-aplikasi atau per-protokol (deep packet inspection)',
    ]
    for item in limitations:
        bullet(doc, '', item)
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # BAB XI — SKENARIO MASALAH & PENANGANAN
    # ═══════════════════════════════════════════
    heading(doc, 'BAB XI — SKENARIO MASALAH & PENANGANANNYA', level=1)
    body(doc,
        'Berikut adalah skenario masalah nyata yang dapat terjadi pada sistem NetVision Pro, '
        'beserta gejala yang terlihat di dashboard, penyebab teknis, dan langkah penanganan.'
    )

    cases = [
        ('Kasus 1: Rx/Tx Speed Selalu 0 Mbps',
         'Dashboard menunjukkan 0.0 Mbps untuk Download dan Upload secara terus-menerus meskipun '
         'ada aktivitas internet di jaringan.',
         [
             'TotalBytesCount Stagnant: Nilai TotalBytesCount:0 di router tidak berubah antar polling. '
             'Ini terjadi karena router ZTE F670L hanya memperbarui counter jika tombol refresh '
             '(#WLANStatus_Btn_refresh) berhasil diklik dan AJAX selesai diproses.',
             'SSID Salah Index: Jika router memiliki multiple SSID (2.4G + 5G), traffic mungkin '
             'tercatat di TotalBytesCount:4 bukan :0. Sistem sudah mengiterasi :0 hingga :7 untuk '
             'mengakumulasi semua SSID.',
             'Siklus Pertama: Kecepatan selalu 0 pada polling pertama karena tidak ada data '
             'sebelumnya (prevState.timestamp = 0) untuk dihitung selisihnya.',
             'Counter Reset Router: Jika router di-restart, rx_bytes baru lebih kecil dari '
             'rx_bytes lama → rxDiff < 0 → sistem mengset kecepatan ke 0 (guard overflow).',
         ],
         'Cek log terminal backend. Jika tampil "Scrape sukses → Rx: XXXXXX bytes" tapi angkanya '
         'tidak berubah → masalah di router (coba refresh manual halaman admin router). '
         'Jika angka berubah tapi dashboard tetap 0 → periksa kalkulasi delta di server.js.'),

        ('Kasus 2: Anomali Terdeteksi Terus-Menerus (False Alarm)',
         'Dashboard menunjukkan status ANOMALI dan Network Score Grade D secara terus-menerus '
         'meskipun jaringan terasa normal.',
         [
             'Model Belum Dilatih Cukup: Jika database memiliki < 20 record, sistem menggunakan '
             'threshold statis (CPU > 85%). Jika CPU server memang sering di atas 85%, ini '
             'akan selalu memicu anomali.',
             'Distribusi Data Tidak Normal: Jika 100 log historis mencakup periode anomali '
             'yang banyak, mean (μ) terangkat tinggi dan stdDev (σ) membesar, menurunkan '
             'sensitivitas model.',
             'CPU Server Tinggi Akibat Puppeteer: Headless Chrome mengkonsumsi CPU signifikan '
             'saat scraping. Jika CPU server memang tinggi dan melebihi 60% + Z > 3, '
             'anomali sah dipicu meski jaringan normal.',
         ],
         'Cek /api/model/stats untuk melihat nilai μ dan σ terkini. Jika σ sangat kecil '
         '(mendekati 0) dan CPU saat ini sedikit lebih tinggi dari μ, Z-Score akan besar. '
         'Solusi: naikkan ANOMALY_THRESHOLD di .env (contoh: 4.0) atau tunggu model '
         'terakumulasi lebih banyak data normal.'),

        ('Kasus 3: Login Router Gagal (Error Credentials)',
         'Log terminal menampilkan "Login gagal! Periksa ROUTER_USER dan ROUTER_PASS di file .env" '
         'secara berulang. Dashboard menampilkan status OFFLINE.',
         [
             'Password Router Berubah: ISP atau pengguna mengubah password router secara manual '
             'melalui halaman admin ZTE F670L.',
             'Selector Login Form Berbeda: Firmware update ZTE F670L mengubah ID elemen form '
             'login (contoh: #Frm_Username berubah menjadi #Username).',
             'Session Conflict: Jika ada pengguna lain yang login ke admin router secara '
             'bersamaan, session Puppeteer bisa diinvalidasi.',
         ],
         'Update nilai ROUTER_PASS di file .env sesuai password router terbaru. '
         'Restart backend setelah perubahan. Jika selector berubah, buka debug_router.js '
         'dan sesuaikan selector form login.'),

        ('Kasus 4: Dashboard Menampilkan 403 Forbidden saat Refresh',
         'Saat pengguna menekan F5 (refresh browser) pada URL selain root (/), browser '
         'mendapatkan respons 403 Forbidden dari server.',
         [
             'SPA Routing Issue: Karena NetVision Pro adalah Single Page Application (SPA), '
             'semua routing ditangani oleh React di sisi klien. Saat browser me-refresh URL '
             'seperti /dashboard atau /stats, Express mencoba mencari file statis tersebut '
             'di dist/ yang tidak ada, mengembalikan 403.',
         ],
         'Solusi sudah diimplementasi: wildcard catch-all route "app.get(\'*\', ...)" di server.js '
         'mengarahkan semua request yang tidak cocok ke index.html, membiarkan React Router '
         'menangani navigasi. Pastikan baris ini ada dan berada setelah semua route API.'),

        ('Kasus 5: Klien WiFi Tidak Muncul di Dashboard',
         'Widget "Client Terkoneksi" menampilkan "Tidak ada client WiFi terdeteksi" meskipun '
         'ada beberapa perangkat yang jelas-jelas terhubung ke WiFi.',
         [
             'Panel Wlan_ClientStatBar Tidak Terbuka: Puppeteer gagal mengklik elemen '
             '#Wlan_ClientStatBar (mungkin elemen tidak ada atau sudah berubah nama).',
             'Refresh Client Timeout: Tombol #Btn_refresh_Wlan_ClientStat tidak ditemukan '
             'atau request AJAX untuk memuat daftar klien belum selesai saat Puppeteer '
             'mencoba membaca elemen klien.',
             'Struktur HTML Berbeda: Template elemen klien berubah dari '
             '"template_Wlan_ClientStat_N" ke format ID yang berbeda setelah firmware update.',
         ],
         'Jalankan debug_router.js untuk mengambil screenshot halaman router dan menganalisis '
         'struktur HTML aktual. Sesuaikan selector di fungsi scrapeRouterStats() di server.js.'),

        ('Kasus 6: Browser Crash / Puppeteer Error',
         'Log terminal menampilkan "Crash terdeteksi. Meluncurkan ulang browser..." atau '
         '"TargetCloseError: Target closed". Dashboard menunjukkan data lama atau OFFLINE.',
         [
             'Memory Leak Headless Chrome: Setelah berjalan lama (beberapa jam/hari), '
             'headless Chrome dapat mengalami kebocoran memori yang menyebabkan crash.',
             'Resource Exhaustion: Sistem operasi kehabisan RAM atau handle file, '
             'menyebabkan proses Chrome tidak dapat dilanjutkan.',
             'Network Timeout: Koneksi ke router timeout lebih dari 15 detik (timeout default '
             'Puppeteer navigation), menyebabkan operasi dibatalkan dan browser crash.',
         ],
         'Sistem sudah memiliki recovery otomatis: getBrowserInstance() mendeteksi crash '
         'dan meluncurkan ulang browser pada siklus berikutnya. Jika terjadi terus-menerus, '
         'pertimbangkan menambahkan memory limit atau menggunakan puppeteer.connect() '
         'ke Chrome yang sudah berjalan terpisah.'),

        ('Kasus 7: Data Supabase Tidak Tersimpan',
         'Log terminal menampilkan "[DB ERROR]" atau tidak ada entri "[DB] ✅ Data tersimpan". '
         'Widget Log Database Historis kosong atau tidak diperbarui.',
         [
             'Kredensial Supabase Kedaluwarsa: API key Supabase di .env sudah tidak valid '
             'atau project Supabase di-pause (versi gratis Supabase meng-pause proyek '
             'yang tidak aktif selama lebih dari 7 hari).',
             'RLS Aktif: Row-Level Security di tabel network_logs mengblokir operasi INSERT '
             'dari key anon. Perlu policy "Allow all" atau service_role key.',
             'Schema Mismatch: Kolom yang dikirim backend (termasuk client_count) tidak ada '
             'di tabel Supabase karena belum ditambahkan.',
         ],
         'Buka dashboard Supabase → pastikan project aktif (tidak di-pause). Periksa RLS '
         'di Table Editor → Policies. Pastikan kolom client_count sudah ditambahkan dengan '
         'tipe int4, nullable true.'),

        ('Kasus 8: Cloudflare Tunnel Tidak Terhubung',
         'URL publik https://jarkom.clyuti.my.id menampilkan halaman error Cloudflare '
         '(502 Bad Gateway atau "Tunnel Not Found").',
         [
             'Proses cloudflared Mati: Proses background Cloudflare Tunnel berhenti '
             '(mungkin karena restart komputer).',
             'Backend Offline: Cloudflare Tunnel berhasil terhubung ke Cloudflare, tapi '
             'tidak ada server yang mendengarkan di localhost:5000.',
             'Konfigurasi Tunnel: File konfigurasi tunnel (credentials) korup atau '
             'token sudah tidak valid.',
         ],
         'Jalankan ulang: "& cloudflared.exe tunnel run jarkom". Pastikan backend juga '
         'berjalan di port 5000 sebelum tunnel diaktifkan. Cek health check lokal '
         'di http://localhost:5000/api/health terlebih dahulu.'),
    ]

    for case_title, symptom, causes, solution in cases:
        heading(doc, case_title, level=2)
        p = doc.add_paragraph()
        set_para_spacing(p, 0, 4)
        add_run(p, '🔴 Gejala: ', bold=True, color=(220,38,38))
        add_run(p, symptom)

        p2 = doc.add_paragraph()
        set_para_spacing(p2, 4, 4)
        add_run(p2, '🔍 Kemungkinan Penyebab:', bold=True, color=(71,85,105))
        for cause in causes:
            bullet(doc, '', cause)

        p3 = doc.add_paragraph()
        set_para_spacing(p3, 4, 10)
        add_run(p3, '✅ Langkah Penanganan: ', bold=True, color=(5,150,105))
        add_run(p3, solution)
        divider(doc)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # BAB XII — KELEBIHAN & KEKURANGAN
    # ═══════════════════════════════════════════
    heading(doc, 'BAB XII — KELEBIHAN & KEKURANGAN SISTEM', level=1)

    heading(doc, '12.1 Kelebihan Sistem', level=2)
    strengths = [
        ('Deteksi Anomali Adaptif (Self-Learning): ',
         'Model Z-Score yang dilatih ulang setiap 5 menit membuat sistem lebih akurat dari '
         'threshold statis. Batas deteksi otomatis menyesuaikan pola penggunaan harian '
         '(ramai siang, sepi malam), drastis mengurangi false alarm.'),
        ('Tidak Bergantung pada SNMP: ',
         'Menggunakan web scraping Puppeteer sebagai alternatif cerdas untuk router '
         'consumer-grade yang tidak menyediakan SNMP atau API. Metode ini berlaku untuk '
         'berbagai router rumahan yang memiliki antarmuka admin web.'),
        ('Page Singleton Pattern (Efisien): ',
         'Menggunakan satu tab browser yang persisten (tidak menutup/membuka tab baru '
         'setiap siklus) menurunkan overhead scraping dari 15–20 detik menjadi 3–5 detik.'),
        ('Arsitektur Full-Stack Terpadu: ',
         'Satu perintah "npm start" meluncurkan backend yang juga melayani frontend statis. '
         'Tidak diperlukan konfigurasi server terpisah atau reverse proxy kompleks.'),
        ('Multi-Dimensi Monitoring: ',
         'Memantau CPU, download speed, upload speed, uptime, dan klien WiFi sekaligus — '
         'memberikan gambaran holistik kondisi jaringan dan server dalam satu layar.'),
        ('Aksesibilitas Publik via Tunnel: ',
         'Cloudflare Tunnel memungkinkan dashboard diakses dari mana saja tanpa port '
         'forwarding, sangat praktis untuk demonstrasi presentasi di kampus.'),
        ('Transparansi Model (Model Stats API): ',
         'Endpoint /api/model/stats mengekspos statistik model secara real-time (μ, σ, Z-Score '
         'aktif), sangat berguna untuk verifikasi ilmiah dan presentasi akademis.'),
        ('Mock Mode untuk Demo: ',
         'MOCK_MODE=true menghasilkan data simulasi dengan anomali terjadwal setiap ~2 menit, '
         'memungkinkan demonstrasi dashboard tanpa koneksi ke router fisik.'),
        ('Pelacakan Active Since Klien: ',
         'Map clientFirstSeen di memori backend mencatat waktu pertama setiap perangkat '
         'terdeteksi, memudahkan identifikasi perangkat baru atau tidak dikenal.'),
        ('Network Score yang Intuitif: ',
         'Mengkonversi data teknis kompleks menjadi nilai 0–100 dan grade A/B/C/D yang '
         'mudah dipahami oleh pengguna non-teknis.'),
    ]
    for label, desc in strengths:
        bullet(doc, label, desc)

    heading(doc, '12.2 Kekurangan & Batasan Sistem', level=2)
    weaknesses = [
        ('Ketergantungan pada Struktur HTML Router: ',
         'Web scraping sangat sensitif terhadap perubahan antarmuka web router. Firmware '
         'update router ZTE F670L dapat mengubah ID elemen HTML, menyebabkan scraping '
         'gagal total tanpa modifikasi kode. Tidak ada mekanisme auto-detect.'),
        ('Tidak Ada Monitoring Per-Client Bandwidth: ',
         'Sistem hanya dapat membaca total bandwidth seluruh jaringan, bukan usage '
         'per-perangkat. Identifikasi perangkat mana yang mengkonsumsi bandwidth berlebih '
         'tidak dimungkinkan dengan pendekatan scraping saat ini.'),
        ('Active Since Tidak Persisten: ',
         'Data waktu "active since" klien hilang setiap kali backend di-restart karena '
         'disimpan hanya di memori (Map), bukan di database. Klien yang online sebelum '
         'restart akan muncul sebagai "baru" setelah restart.'),
        ('Beban CPU oleh Puppeteer: ',
         'Headless Chrome memerlukan RAM 100–300 MB dan meningkatkan CPU penggunaan '
         'saat scraping. Pada komputer dengan spesifikasi rendah, ini dapat '
         'mempengaruhi akurasi metrik CPU yang dipantau.'),
        ('Single Point of Failure: ',
         'Seluruh sistem berjalan di satu komputer lokal. Jika komputer mati, '
         'monitoring berhenti. Tidak ada failover, redundansi, atau notifikasi '
         'push (email/WhatsApp) ketika anomali terdeteksi.'),
        ('Keterbatasan Gratis Supabase: ',
         'Versi gratis Supabase meng-pause proyek yang tidak aktif selama 7 hari. '
         'Penyimpanan dibatasi 500MB. Untuk penggunaan produksi jangka panjang, '
         'diperlukan paket berbayar atau self-hosted PostgreSQL.'),
        ('Tidak Ada Autentikasi Dashboard: ',
         'Dashboard dapat diakses oleh siapa saja yang mengetahui URL publik '
         'tanpa login. Informasi sensitif seperti IP router, MAC Address, dan '
         'nama perangkat terekspos ke internet publik.'),
        ('Hanya Kompatibel dengan Windows: ',
         'Pengambilan data CPU dan Uptime menggunakan perintah PowerShell '
         '(Get-CimInstance Win32_Processor) yang hanya tersedia di Windows. '
         'Tidak dapat berjalan di Linux/macOS tanpa modifikasi kode.'),
        ('Interval Polling Tetap (30 Detik): ',
         'Tidak ada mekanisme adaptive polling — polling tetap setiap 30 detik '
         'meskipun jaringan dalam kondisi kritis dan membutuhkan update lebih cepat. '
         'Polling lebih cepat berisiko memicu rate limit atau overloading router.'),
        ('Tidak Ada Notifikasi Aktif: ',
         'Anomali hanya ditampilkan secara visual di dashboard. Tidak ada notifikasi '
         'push (email, Telegram bot, atau SMS) yang dikirim saat anomali pertama '
         'kali terdeteksi. Pengguna harus aktif membuka dashboard.'),
    ]
    for label, desc in weaknesses:
        bullet(doc, label, desc)
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # BAB XIII — KESIMPULAN
    # ═══════════════════════════════════════════
    heading(doc, 'BAB XIII — KESIMPULAN', level=1)

    body(doc,
        'NetVision Pro adalah solusi monitoring jaringan WiFi yang inovatif untuk lingkungan '
        'rumahan dan skala kecil, terutama di mana router yang digunakan (seperti ZTE F670L '
        'dari IndiHome) tidak menyediakan protokol monitoring standar seperti SNMP. '
        'Dengan menggabungkan teknik web scraping headless, algoritma deteksi anomali statistik '
        'adaptif (Z-Score), dan antarmuka dashboard web modern, aplikasi ini berhasil '
        'mengisi kesenjangan antara kebutuhan monitoring profesional dengan keterbatasan '
        'hardware consumer-grade.'
    )
    body(doc,
        'Implementasi Page Singleton Pattern pada Puppeteer terbukti secara signifikan '
        'meningkatkan efisiensi scraping (dari 15–20 detik menjadi 3–5 detik per siklus), '
        'sementara mekanisme Dynamic Re-training setiap 5 menit memastikan deteksi anomali '
        'tetap akurat mengikuti perubahan pola penggunaan jaringan harian. Dual-guard condition '
        'pada algoritma Z-Score (Z-Score > 3.0 DAN nilai fisik minimum) berhasil mengurangi '
        'false alarm yang menjadi kelemahan utama sistem threshold statis konvensional.'
    )
    body(doc,
        'Kekurangan utama sistem ini — ketergantungan pada struktur HTML router dan '
        'ketiadaan monitoring per-client bandwidth — membuka peluang pengembangan '
        'lanjutan. Integrasi dengan protokol SSH router, penambahan notifikasi push, '
        'autentikasi dashboard, dan portabilitas ke Linux/macOS adalah area '
        'pengembangan yang dapat memperluas nilai dan keandalan aplikasi ini secara '
        'signifikan untuk digunakan dalam lingkungan yang lebih luas.'
    )
    body(doc,
        'Secara keseluruhan, NetVision Pro berhasil mencapai tujuan utamanya: '
        'memberikan visibilitas real-time kondisi jaringan nirkabel beserta deteksi '
        'otomatis perilaku abnormal, yang dapat diakses dari mana saja melalui URL publik, '
        'dengan arsitektur yang efisien dan mudah dikonfigurasi.'
    )

    divider(doc)
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_end.paragraph_format.space_before = Pt(20)
    add_run(p_end,
            f'Dokumen dibuat secara otomatis oleh generate_full_doc.py\n'
            f'NetVision Pro — Kelompok 6 Jaringan Komputer — {datetime.date.today().strftime("%d %B %Y")}',
            italic=True, color=(148,163,184), size=9)

    doc.save(OUTPUT_PATH)
    print(f'\n✅ Dokumen berhasil disimpan ke:\n   {OUTPUT_PATH}')

if __name__ == '__main__':
    build()
