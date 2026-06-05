# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime

OUTPUT_PATH = r"d:\jarkom\Catatan_Presentasi_NetVision_Pro.docx"

# ─── HELPERS ───────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc.append(shd)

def set_cell_padding(cell, top=100, bottom=100, left=160, right=160):
    tc = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for name, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        node = OxmlElement(f'w:{name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tc.append(tcMar)

def add_borders(table, color='CBD5E1'):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tblPr.append(borders)

def run(para, text, bold=False, italic=False, size=11,
        color=(30,41,59), font='Arial'):
    r = para.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = font; r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*color)
    return r

def heading1(doc, text, before=16, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.keep_with_next = True
    run(p, text, bold=True, size=17, color=(15,23,42))
    return p

def heading2(doc, text, before=12, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.keep_with_next = True
    run(p, text, bold=True, size=13, color=(37,99,235))
    return p

def heading3(doc, text, before=8, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.keep_with_next = True
    run(p, text, bold=True, size=11, color=(71,85,105))
    return p

def body(doc, text, after=7, italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run(p, text, italic=italic)
    return p

def bullet(doc, label, desc, after=5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(after)
    if label:
        run(p, label, bold=True, color=(37,99,235))
    run(p, desc)
    return p

def tip_box(doc, title, content, bg='EFF6FF', border='BFDBFE', title_color=(30,64,175)):
    """Styled callout box using a 1-cell table."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, bg)
    set_cell_padding(cell, 160, 160, 200, 200)
    cell.paragraphs[0].clear()
    # Border
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '12')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), border)
        tcBorders.append(b)
    tcPr.append(tcBorders)
    # Title
    p_title = cell.paragraphs[0]
    r_t = p_title.add_run(title)
    r_t.bold = True; r_t.font.name = 'Arial'
    r_t.font.size = Pt(10.5)
    r_t.font.color.rgb = RGBColor(*title_color)
    # Content
    for line in content:
        p_c = cell.add_paragraph()
        r_c = p_c.add_run(line)
        r_c.font.name = 'Arial'; r_c.font.size = Pt(10)
        r_c.font.color.rgb = RGBColor(30,41,59)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)
    return t

def qa_box(doc, question, answer):
    """Q&A styled box."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, 'FFF7ED')
    set_cell_padding(cell, 160, 160, 200, 200)
    tcPr = cell._tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '12')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'FED7AA')
        tcB.append(b)
    tcPr.append(tcB)
    cell.paragraphs[0].clear()
    p_q = cell.paragraphs[0]
    r_q = p_q.add_run('T: ' + question)
    r_q.bold = True; r_q.font.name = 'Arial'
    r_q.font.size = Pt(10.5)
    r_q.font.color.rgb = RGBColor(180,83,9)
    p_a = cell.add_paragraph()
    r_a = p_a.add_run('J: ' + answer)
    r_a.font.name = 'Arial'; r_a.font.size = Pt(10)
    r_a.font.color.rgb = RGBColor(30,41,59)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)
    return t

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run('─' * 95)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(203,213,225)

def two_col(doc, rows, header_bg='1E3A5F', headers=None):
    total = len(rows) + (1 if headers else 0)
    t = doc.add_table(rows=total, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_borders(t)
    idx = 0
    if headers:
        hcells = t.rows[0].cells
        for i,h in enumerate(headers):
            set_cell_bg(hcells[i], header_bg)
            set_cell_padding(hcells[i], 120, 120, 160, 160)
            r = hcells[i].paragraphs[0].add_run(h)
            r.bold=True; r.font.size=Pt(10)
            r.font.color.rgb=RGBColor(255,255,255)
        idx = 1
    for label, val in rows:
        cells = t.rows[idx].cells
        set_cell_bg(cells[0], 'F8FAFC')
        set_cell_padding(cells[0], 100, 100, 160, 160)
        set_cell_padding(cells[1], 100, 100, 160, 160)
        r0 = cells[0].paragraphs[0].add_run(label)
        r0.bold=True; r0.font.size=Pt(10)
        r0.font.color.rgb=RGBColor(51,65,85)
        r1 = cells[1].paragraphs[0].add_run(val)
        r1.font.size=Pt(10)
        r1.font.color.rgb=RGBColor(30,41,59)
        idx += 1
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t

# ─── DOCUMENT BUILDER ──────────────────────────────────

def build():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.15)
        sec.right_margin  = Inches(1.0)

    # ══════════════════════════════════════════
    # COVER
    # ══════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run('CATATAN PRESENTASI')
    r.bold=True; r.font.name='Arial'; r.font.size=Pt(14)
    r.font.color.rgb=RGBColor(71,85,105)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run('NetVision Pro')
    r2.bold=True; r2.font.name='Arial'; r2.font.size=Pt(30)
    r2.font.color.rgb=RGBColor(37,99,235)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(40)
    r3 = p3.add_run('Network Monitoring Dashboard & Anomaly Detection System')
    r3.italic=True; r3.font.name='Arial'; r3.font.size=Pt(13)
    r3.font.color.rgb=RGBColor(100,116,139)

    divider(doc)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(16)
    r4 = p4.add_run(
        f'Kelompok 6  |  Jaringan Komputer  |  {datetime.date.today().strftime("%d %B %Y")}'
    )
    r4.font.name='Arial'; r4.font.size=Pt(11)
    r4.font.color.rgb=RGBColor(100,116,139)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # BAB 1 — PEMBUKA PRESENTASI
    # ══════════════════════════════════════════
    heading1(doc, 'BAGIAN 1 — PEMBUKA: APA ITU NETVISION PRO?')

    tip_box(doc,
        'Script Pembuka (bisa dibaca saat presentasi):',
        [
            '"Kami memperkenalkan NetVision Pro — sebuah sistem monitoring jaringan WiFi berbasis web '
            'yang kami kembangkan untuk memantau performa router ZTE F670L secara real-time. '
            'Berbeda dengan tools monitoring konvensional, aplikasi kami mampu mendeteksi anomali '
            'lalu lintas data secara otomatis menggunakan kecerdasan statistik, '
            'dan dapat diakses dari mana saja melalui internet."'
        ],
        bg='EFF6FF', border='93C5FD', title_color=(30,64,175)
    )

    heading2(doc, '1.1 Definisi Singkat')
    body(doc,
        'NetVision Pro adalah aplikasi web full-stack yang berfungsi sebagai "mata" pengawas '
        'jaringan WiFi Anda. Ia secara otomatis mengumpulkan, memproses, menyimpan, dan '
        'menampilkan data kondisi jaringan secara visual dan interaktif — 24 jam penuh, '
        'tanpa perlu intervensi manual.'
    )

    heading2(doc, '1.2 Masalah yang Diselesaikan')
    body(doc, 'Tanpa NetVision Pro, administrator jaringan rumahan menghadapi tiga masalah utama:')
    bullet(doc, 'Buta terhadap kondisi jaringan: ',
           'Tidak ada cara mudah untuk mengetahui apakah internet sedang lambat karena '
           'bandwidth habis, CPU server penuh, atau ada perangkat asing yang masuk.')
    bullet(doc, 'Tidak ada riwayat data: ',
           'Jika jaringan bermasalah kemarin malam, tidak ada catatan yang bisa diperiksa.')
    bullet(doc, 'Alarm palsu yang mengganggu: ',
           'Sistem monitoring lama menggunakan batas statis (fixed threshold) — misal "anomali '
           'jika upload > 20 Mbps" — yang tidak mempertimbangkan bahwa pola penggunaan jaringan '
           'berubah-ubah setiap hari.')

    heading2(doc, '1.3 Solusi yang Ditawarkan')
    bullet(doc, 'Monitoring Real-Time: ',
           'Dashboard web yang menampilkan kecepatan download/upload, beban CPU, uptime server, '
           'dan daftar perangkat yang terhubung, diperbarui setiap 5 detik.')
    bullet(doc, 'Penyimpanan Historis Permanen: ',
           'Setiap data dicatat ke database cloud Supabase — tersedia kapan saja untuk '
           'analisis mundur, bahkan setelah aplikasi di-restart.')
    bullet(doc, 'Deteksi Anomali Cerdas (Z-Score Dinamis): ',
           'Batas deteksi anomali dihitung ulang otomatis setiap 5 menit berdasarkan pola '
           'historis, sehingga sangat akurat dan bebas alarm palsu.')
    bullet(doc, 'Aksesibilitas Publik: ',
           'Dashboard dapat dibuka dari luar jaringan rumah melalui URL publik '
           'https://jarkom.clyuti.my.id tanpa konfigurasi rumit.')

    doc.add_page_break()

    # ══════════════════════════════════════════
    # BAB 2 — ARSITEKTUR & CARA KERJA
    # ══════════════════════════════════════════
    heading1(doc, 'BAGIAN 2 — ARSITEKTUR & CARA KERJA SISTEM')

    heading2(doc, '2.1 Empat Komponen Utama')
    body(doc,
        'Sistem NetVision Pro terdiri dari empat komponen yang saling bekerja sama. '
        'Bayangkan seperti sebuah pabrik dengan lini produksi:'
    )

    two_col(doc, [
        ('1. Data Collector (Backend Node.js)',
         'Mesin pengambil data otomatis. Setiap 30 detik, ia masuk ke halaman admin '
         'router ZTE F670L, membaca data bandwidth dan daftar perangkat WiFi, '
         'lalu mengambil data CPU dari sistem Windows. Ini adalah "kaki" '
         'yang berjalan terus-menerus di latar belakang.'),
        ('2. Database Cloud (Supabase)',
         'Gudang penyimpanan data permanen di cloud. Setiap data yang berhasil dikumpulkan '
         'langsung disimpan di sini secara otomatis. Data tersedia 24/7 dan '
         'tidak hilang meski aplikasi di-restart.'),
        ('3. Dashboard Web (Frontend React.js)',
         'Wajah aplikasi yang dilihat pengguna. Ia mengambil data dari backend setiap '
         '5 detik dan menampilkannya dalam grafik, angka, dan indikator berwarna '
         'yang mudah dipahami sekilas pandang.'),
        ('4. Cloudflare Tunnel',
         'Jembatan aman yang menghubungkan server lokal ke internet publik. '
         'Berkat ini, dashboard dapat diakses dari kampus, rumah teman, atau '
         'mana saja — tanpa perlu konfigurasi router ISP.'),
    ], headers=['Komponen', 'Peran & Fungsi'])

    heading2(doc, '2.2 Alur Data dari Ujung ke Ujung')
    body(doc,
        'Berikut adalah perjalanan data dari sumber (router) hingga tampil di layar '
        'dalam waktu kurang dari 5 detik:'
    )

    steps = [
        ('Langkah 1 — Puppeteer Login ke Router',
         'Sebuah browser Chrome tak kasat mata (headless) secara otomatis membuka '
         'http://192.168.1.1, mengisi username/password, dan masuk ke halaman '
         'admin router ZTE F670L.'),
        ('Langkah 2 — Scraping Data Bandwidth',
         'Puppeteer menavigasi ke menu "WLAN Status" dan membaca nilai elemen '
         'TotalBytesCount yang berisi akumulasi total byte yang dikirim dan diterima '
         'router sejak menyala. Format datanya: "rxBytes/txBytes".'),
        ('Langkah 3 — Scraping Daftar Client WiFi',
         'Puppeteer juga membuka panel "WLAN Client Status" untuk membaca siapa '
         'saja perangkat yang sedang terhubung — nama perangkat, IP, MAC address, '
         'dan nama WiFi (SSID) yang digunakan.'),
        ('Langkah 4 — Ambil Data CPU & Uptime via PowerShell',
         'Bersamaan dengan scraping router, sistem menjalankan perintah PowerShell '
         'untuk membaca beban prosesor (CPU%) dan berapa lama komputer sudah '
         'menyala (uptime) dari sistem operasi Windows.'),
        ('Langkah 5 — Kalkulasi Kecepatan Mbps',
         'Karena router hanya menyimpan total byte kumulatif, backend menghitung '
         'kecepatan dengan rumus: Kecepatan = (Byte Sekarang - Byte 30 detik lalu) '
         'x 8 dibagi (30 x 1.000.000). Hasilnya: Mbps aktual saat ini.'),
        ('Langkah 6 — Deteksi Anomali Z-Score',
         'Nilai CPU, download, dan upload dibandingkan dengan model statistik. '
         'Jika nilainya menyimpang lebih dari 3 kali standar deviasi dari rata-rata '
         'historis, data ditandai sebagai anomali.'),
        ('Langkah 7 — Simpan ke Supabase',
         'Semua data (CPU, kecepatan, jumlah client, status anomali) dikirim ke '
         'tabel network_logs di Supabase Cloud PostgreSQL. Data ini tersimpan '
         'permanen dan bisa diakses kapan saja.'),
        ('Langkah 8 — Frontend Ambil & Tampilkan',
         'React dashboard setiap 5 detik mengambil data terbaru dari backend, '
         'lalu memperbarui semua widget — angka kecepatan, grafik, daftar client, '
         'skor jaringan — secara otomatis tanpa perlu refresh halaman.'),
    ]
    for title, desc in steps:
        heading3(doc, title)
        body(doc, desc, indent=True)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # BAB 3 — SUPABASE (highlight khusus)
    # ══════════════════════════════════════════
    heading1(doc, 'BAGIAN 3 — MENGAPA SUPABASE? (PENJELASAN MENDALAM)')

    tip_box(doc,
        'Poin Unggulan untuk Presentasi:',
        [
            'Supabase adalah salah satu keputusan arsitektur TERBAIK dalam proyek ini. '
            'Jelaskan dengan antusias — ini yang membuat sistem jauh lebih powerful '
            'dibandingkan monitoring sederhana yang hanya menyimpan data di RAM.',
        ],
        bg='F0FDF4', border='86EFAC', title_color=(5,122,85)
    )

    heading2(doc, '3.1 Apa Itu Supabase?')
    body(doc,
        'Supabase adalah platform database cloud open-source yang menyediakan '
        'PostgreSQL (database relasional yang andal) beserta API otomatis, '
        'autentikasi, dan storage — semuanya gratis untuk proyek skala kecil. '
        'Ia sering disebut sebagai "Firebase alternatif open-source".'
    )
    body(doc,
        'Dalam proyek ini, Supabase berperan sebagai gudang data permanen '
        'untuk semua log monitoring jaringan yang dikumpulkan oleh NetVision Pro.'
    )

    heading2(doc, '3.2 Kenapa Tidak Simpan di RAM Saja?')
    body(doc,
        'Pertanyaan yang wajar: "Mengapa perlu database? Bukankah bisa '
        'simpan data di variabel JavaScript saja?" Jawabannya adalah:'
    )

    two_col(doc, [
        ('Data di RAM (tanpa Supabase)',
         'Hilang setiap kali backend di-restart atau komputer mati. '
         'Grafik selalu mulai dari nol. Tidak bisa analisis data kemarin. '
         'Tidak bisa akses dari perangkat lain.'),
        ('Data di Supabase (cloud)',
         'Tersimpan permanen walaupun backend di-restart 100 kali. '
         'Grafik selalu menampilkan 50 data terakhir dari semua sesi sebelumnya. '
         'Data bisa dianalisis kapan saja dan dari mana saja.'),
    ], headers=['Pendekatan', 'Konsekuensi'])

    heading2(doc, '3.3 Mengapa Data Jam 11 Muncul Padahal Baru Dijalankan Jam 2?')
    body(doc,
        'Ini adalah pertanyaan yang sangat bagus dan menunjukkan betapa bergunanya Supabase. '
        'Berikut penjelasannya:'
    )
    tip_box(doc,
        'Penjelasan Singkat untuk Audiens:',
        [
            '"Perhatikan bahwa meskipun kami baru menjalankan ulang aplikasi jam 2 siang, '
            'grafik masih menampilkan data dari jam 11 pagi. Ini karena setiap data '
            'yang dikumpulkan disimpan secara permanen di Supabase Cloud. Ketika '
            'dashboard dibuka, ia langsung mengambil 50 data terakhir dari database '
            '— termasuk data dari sesi-sesi sebelumnya. Data tidak pernah hilang, '
            'bahkan setelah restart berkali-kali. Inilah keunggulan menggunakan '
            'database cloud dibandingkan penyimpanan di memori biasa."'
        ],
        bg='EFF6FF', border='93C5FD', title_color=(30,64,175)
    )

    body(doc, 'Alur lengkapnya:')
    bullet(doc, 'Sesi pagi (jam 11): ',
           'Backend berjalan, scraping router, menyimpan data ke Supabase setiap 30 detik.')
    bullet(doc, 'Backend restart (jam 14): ',
           'Semua variabel di memori (RAM) hilang. Tapi data di Supabase tetap aman.')
    bullet(doc, 'Dashboard dibuka (jam 14:20): ',
           'Frontend meminta data ke endpoint /api/metrics/history. Backend '
           'mengambil 50 record TERBARU dari Supabase — yang mencakup data dari '
           'jam 11 tadi. Semua tampil di grafik secara otomatis.')

    heading2(doc, '3.4 Struktur Tabel network_logs di Supabase')
    body(doc,
        'Supabase menyimpan data monitoring dalam satu tabel bernama network_logs. '
        'Setiap 30 detik, satu baris data baru ditambahkan:'
    )
    two_col(doc, [
        ('id',             'Nomor unik otomatis untuk setiap record'),
        ('created_at',     'Waktu pencatatan (timestamp otomatis dari Supabase)'),
        ('cpu_usage',      'Beban CPU komputer server (contoh: 45)'),
        ('uptime',         'Berapa lama server menyala (contoh: 258h 49m 42s)'),
        ('rx_bytes',       'Total byte yang diterima router sejak menyala'),
        ('tx_bytes',       'Total byte yang dikirim router sejak menyala'),
        ('client_count',   'Jumlah perangkat WiFi yang terhubung saat itu'),
        ('is_anomaly',     'True jika terdeteksi anomali, False jika normal'),
        ('anomaly_reason', 'Kalimat penjelasan kenapa dianggap anomali'),
        ('router_ip',      'IP address router yang dimonitor'),
    ], headers=['Nama Kolom', 'Isi & Fungsi'])

    heading2(doc, '3.5 Keunggulan Teknis Supabase dalam Proyek Ini')
    bullet(doc, 'Gratis & Tanpa Batas Record: ',
           'Paket gratis Supabase sudah cukup untuk menyimpan ribuan log monitoring '
           'dengan storage hingga 500 MB — lebih dari cukup untuk proyek ini.')
    bullet(doc, 'API Otomatis: ',
           'Supabase JS Client memungkinkan backend menyimpan dan mengambil data '
           'hanya dengan beberapa baris kode, tanpa perlu menulis query SQL kompleks.')
    bullet(doc, 'Realtime Capability: ',
           'Supabase mendukung subscription realtime — jika diperlukan di masa depan, '
           'frontend bisa langsung menerima notifikasi saat data baru masuk tanpa polling.')
    bullet(doc, 'Dashboard Admin: ',
           'Tim bisa melihat seluruh data log langsung di dashboard web Supabase '
           '(app.supabase.com) kapan saja untuk keperluan verifikasi atau debug.')

    doc.add_page_break()

    # ══════════════════════════════════════════
    # BAB 4 — DETEKSI ANOMALI Z-SCORE
    # ══════════════════════════════════════════
    heading1(doc, 'BAGIAN 4 — DETEKSI ANOMALI: ALGORITMA Z-SCORE DINAMIS')

    tip_box(doc,
        'Poin Kunci untuk Presentasi:',
        [
            'Ini adalah fitur yang paling membedakan NetVision Pro dari monitoring biasa. '
            'Tekankan kata "DINAMIS" dan "SELF-LEARNING" — model ini belajar sendiri '
            'dari data historis dan menyesuaikan batasnya setiap 5 menit.'
        ],
        bg='F0FDF4', border='86EFAC', title_color=(5,122,85)
    )

    heading2(doc, '4.1 Masalah dengan Threshold Statis')
    body(doc,
        'Sistem monitoring lama menggunakan aturan kaku seperti "anomali jika upload > 20 Mbps". '
        'Masalahnya:'
    )
    bullet(doc, 'Jam 3 pagi — jaringan sepi: ',
           'Upload 5 Mbps sudah terasa luar biasa tinggi, tapi sistem tidak mendeteksinya '
           'karena masih di bawah batas 20 Mbps. Anomali TERLEWAT.')
    bullet(doc, 'Jam 8 malam — jam sibuk: ',
           'Upload 18 Mbps adalah hal biasa karena semua orang streaming. Tapi sistem '
           'hampir membunyikan alarm. FALSE ALARM hampir terjadi.')
    body(doc,
        'Kesimpulan: threshold statis tidak bisa membedakan "tinggi yang wajar" '
        'dengan "tinggi yang berbahaya" karena konteks waktu berbeda.',
        italic=True
    )

    heading2(doc, '4.2 Solusi: Z-Score Dinamis')
    body(doc,
        'Z-Score mengukur seberapa jauh suatu nilai menyimpang dari rata-rata '
        'dalam satuan standar deviasi. Rumusnya:'
    )
    p_f = doc.add_paragraph()
    p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_f.paragraph_format.space_before = Pt(6)
    p_f.paragraph_format.space_after  = Pt(6)
    rf = p_f.add_run('Z  =  | X  -  mu |  /  sigma')
    rf.bold = True; rf.font.name = 'Consolas'
    rf.font.size = Pt(16)
    rf.font.color.rgb = RGBColor(37,99,235)

    two_col(doc, [
        ('X',       'Nilai metrik yang baru masuk (Upload saat ini dalam Mbps)'),
        ('mu',      'Rata-rata (mean) dari 100 data historis terakhir di Supabase'),
        ('sigma',   'Standar deviasi dari 100 data historis tersebut'),
        ('Z > 3.0', 'Dianggap anomali — nilainya terlalu jauh dari pola normal'),
    ])

    heading2(doc, '4.3 Contoh Nyata: Upload 33 Mbps Terdeteksi Anomali')
    body(doc,
        'Pada grafik dashboard, terlihat titik anomali merah saat Upload mencapai 33 Mbps. '
        'Berikut perhitungannya dengan data model yang sedang aktif:'
    )
    tip_box(doc,
        'Perhitungan Z-Score Upload (Tx):',
        [
            'mu Tx  = 7.05 Mbps  (rata-rata upload dari 100 log terakhir)',
            'sigma  = 4.56 Mbps  (standar deviasi upload)',
            '',
            'Z = | 33 - 7.05 | / 4.56',
            'Z = 25.95 / 4.56',
            'Z = 5.69',
            '',
            'Karena Z (5.69) > Threshold (3.0), dan nilai fisik Tx (33) > batas minimum (20 Mbps)',
            '=> STATUS: ANOMALI terdeteksi! Titik merah muncul di grafik.',
        ],
        bg='FEF2F2', border='FCA5A5', title_color=(185,28,28)
    )

    heading2(doc, '4.4 Self-Learning: Re-Training Setiap 5 Menit')
    body(doc,
        'Yang membuat sistem ini "pintar" adalah proses pelatihan ulang otomatis '
        'setiap 5 menit. Backend mengambil 100 log terbaru dari Supabase '
        'dan menghitung ulang rata-rata dan standar deviasi untuk setiap metrik. '
        'Artinya:'
    )
    bullet(doc, 'Jam 3 pagi (traffic rendah): ',
           'mu Upload mungkin hanya 2 Mbps dengan sigma 1 Mbps. Upload 6 Mbps '
           'sudah menghasilkan Z = 4 — terdeteksi anomali dengan benar.')
    bullet(doc, 'Jam 8 malam (traffic tinggi): ',
           'mu Upload sudah 10 Mbps dengan sigma 5 Mbps. Upload 18 Mbps hanya '
           'menghasilkan Z = 1.6 — tidak anomali, karena memang normal di jam itu.')
    body(doc,
        'Model menyesuaikan diri dengan ritme penggunaan jaringan secara otomatis, '
        'tanpa ada yang perlu mengkonfigurasinya secara manual.',
        italic=True
    )

    heading2(doc, '4.5 Dual-Guard: Mencegah False Alarm di Traffic Sangat Rendah')
    body(doc,
        'Ada satu kondisi khusus yang perlu diantisipasi: jika traffic SANGAT rendah '
        '(misalnya mu = 0.001 Mbps, sigma = 0.0005 Mbps), maka upload 0.01 Mbps '
        'pun bisa menghasilkan Z > 3. Untuk mencegah ini, sistem menerapkan '
        '"dual-guard" — anomali hanya dipicu jika KEDUA syarat terpenuhi:'
    )
    two_col(doc, [
        ('CPU Load',
         'Z-Score CPU > 3.0  DAN  Nilai CPU > 60%'),
        ('Download (Rx)',
         'Z-Score Rx > 3.0  DAN  Nilai Rx > 50 Mbps'),
        ('Upload (Tx)',
         'Z-Score Tx > 3.0  DAN  Nilai Tx > 20 Mbps'),
    ], headers=['Metrik', 'Syarat Anomali (Harus Keduanya Terpenuhi)'])

    doc.add_page_break()

    # ══════════════════════════════════════════
    # BAB 5 — TEKNIK SCRAPING
    # ══════════════════════════════════════════
    heading1(doc, 'BAGIAN 5 — TEKNIK PENGAMBILAN DATA: WEB SCRAPING PUPPETEER')

    heading2(doc, '5.1 Mengapa Tidak Pakai SNMP?')
    body(doc,
        'Cara standar untuk monitoring router adalah menggunakan protokol SNMP '
        '(Simple Network Management Protocol). Namun, router ZTE F670L yang '
        'disediakan IndiHome MENGUNCI akses SNMP — protokol ini dinonaktifkan '
        'oleh firmware ISP sehingga tidak bisa digunakan.'
    )
    body(doc,
        'Solusi kami: Web Scraping menggunakan Puppeteer — sebuah library Node.js '
        'yang mengendalikan browser Chrome secara programatik, persis seperti '
        'robot yang mengoperasikan browser untuk kita.'
    )

    heading2(doc, '5.2 Page Singleton Pattern: Optimasi Krusial')
    body(doc,
        'Inovasi teknis paling penting dalam implementasi scraping kami adalah '
        '"Page Singleton Pattern". Tanpa optimasi ini, setiap polling 30 detik '
        'akan membuka browser baru, login ulang, navigasi ulang, dan tutup browser '
        '— membutuhkan 15-20 detik per siklus.'
    )

    two_col(doc, [
        ('Tanpa Page Singleton (cara naif)',
         'Setiap 30 detik: Buka Chrome baru (3s) > Login router (5s) > '
         'Navigasi menu (5s) > Baca data (2s) > Tutup Chrome (2s). '
         'Total: ~17 detik per siklus. Browser dibuka-tutup 2.880 kali per hari.'),
        ('Dengan Page Singleton (implementasi kami)',
         'Pertama kali: Login + navigasi seperti biasa (~15s). '
         'Siklus berikutnya: Hanya navigasi ulang ke halaman yang sama (~3-5s). '
         'Browser tetap terbuka, sesi login tetap aktif. '
         'Efisiensi meningkat 70-80% per siklus.'),
    ], headers=['Pendekatan', 'Cara Kerja'])

    heading2(doc, '5.3 Langkah Scraping (Untuk Audiens Teknis)')
    steps2 = [
        ('1. Buka router di browser:', 'Puppeteer membuka http://192.168.1.1'),
        ('2. Cek status login:', 'Jika ada form password, lakukan login otomatis'),
        ('3. Klik menu "Local Network":', 'Menggunakan selector CSS #localnet'),
        ('4. Klik "Status":', 'Menggunakan selector #localNetStatus'),
        ('5. Klik "WLAN Status":', 'Menggunakan selector #WLANStatusBar'),
        ('6. Klik tombol Refresh:', 'Memaksa router update counter bytes (#WLANStatus_Btn_refresh)'),
        ('7. Baca TotalBytesCount:', 'Mengekstrak nilai dari elemen HTML: rxBytes/txBytes'),
        ('8. Buka panel Client Status:', 'Klik #Wlan_ClientStatBar untuk load daftar perangkat'),
        ('9. Baca data setiap client:', 'Hostname, IP, MAC Address, nama WiFi (SSID)'),
    ]
    two_col(doc, steps2, headers=['Langkah', 'Aksi yang Dilakukan'])

    doc.add_page_break()

    # ══════════════════════════════════════════
    # BAB 6 — FITUR-FITUR DASHBOARD
    # ══════════════════════════════════════════
    heading1(doc, 'BAGIAN 6 — FITUR-FITUR DASHBOARD: PENJELASAN TIAP WIDGET')

    heading2(doc, '6.1 Header Bar — Status Sistem')
    body(doc,
        'Bagian paling atas dashboard menampilkan status sistem secara sekilas. '
        'Indikator berwarna memudahkan identifikasi kondisi langsung:'
    )
    bullet(doc, 'Hijau + berdenyut (LIVE): ',
           'Backend terhubung ke router dan data real-time sedang mengalir.')
    bullet(doc, 'Ungu (MOCK MODE): ',
           'Berjalan dalam mode simulasi — cocok untuk demo tanpa router fisik.')
    bullet(doc, 'Merah (OFFLINE): ',
           'Koneksi ke backend terputus — perlu dicek.')
    bullet(doc, 'Badge "X Client": ',
           'Menampilkan jumlah perangkat WiFi yang terhubung saat ini.')
    bullet(doc, 'Tombol Refresh: ',
           'Memaksa pembaruan data manual tanpa menunggu interval 5 detik.')

    heading2(doc, '6.2 Widget Download Speed (Cyan / Biru Terang)')
    bullet(doc, 'Angka besar: ', 'Kecepatan unduh real-time dalam Mbps.')
    bullet(doc, 'Progress bar: ', 'Visual beban bandwidth dari skala 0-100 Mbps.')
    bullet(doc, 'Total RX: ',
           'Total gigabyte yang telah diunduh router sejak pertama menyala — ini '
           'adalah nilai kumulatif dari router, bukan dari backend kita.')

    heading2(doc, '6.3 Widget Upload Speed (Fuchsia / Magenta)')
    body(doc, 'Sama seperti Download, tapi untuk traffic upload dengan skala 0-50 Mbps.')
    tip_box(doc,
        'Catatan Penting untuk Presentasi:',
        [
            'Upload biasanya LEBIH TINGGI dari download di jaringan yang digunakan untuk '
            'cloud storage, video call, atau streaming. Ini normal dan bukan tanda masalah. '
            'Pada dashboard kita, Upload (garis fuchsia/magenta) memang terlihat lebih '
            'dominan dari Download — ini mencerminkan pola penggunaan nyata jaringan kita.',
        ],
        bg='FFF7ED', border='FED7AA', title_color=(146,64,14)
    )

    heading2(doc, '6.4 Widget CPU Load (Gauge Chart Setengah Lingkaran)')
    body(doc,
        'Menampilkan beban prosesor komputer SERVER (bukan router) dalam bentuk '
        'grafik gauge setengah lingkaran yang dinamis. Label status otomatis berubah:'
    )
    bullet(doc, 'Normal (hijau): ', 'CPU <= 50% — sistem berjalan lancar.')
    bullet(doc, 'Beban Sedang (amber/kuning): ', 'CPU 50-85% — perlu diperhatikan.')
    bullet(doc, 'Kritis (merah): ', 'CPU > 85% — kemungkinan masalah performa.')
    body(doc,
        'Data diambil dari PowerShell: Get-CimInstance Win32_Processor. '
        'Bukan CPU router, melainkan CPU komputer yang menjalankan backend.',
        italic=True
    )

    heading2(doc, '6.5 Grafik Live Traffic (LineChart Besar)')
    body(doc,
        'Grafik garis interaktif yang memvisualisasikan tren kecepatan jaringan '
        'dari 50 data historis terakhir (dari Supabase). '
        'Dua garis berjalan bersamaan:'
    )
    bullet(doc, 'Garis Biru (Download/Rx): ', 'Kecepatan unduh dari waktu ke waktu.')
    bullet(doc, 'Garis Fuchsia (Upload/Tx): ', 'Kecepatan unggah dari waktu ke waktu.')
    bullet(doc, 'Titik Merah Berkedip: ',
           'Menandai momen ketika Z-Score melampaui threshold 3.0 — visualisasi '
           'anomali yang langsung terlihat oleh siapapun tanpa perlu memahami statistik.')
    body(doc,
        'Hover di atas grafik akan menampilkan tooltip dengan detail waktu dan nilai '
        'persis pada titik tersebut.',
        italic=True
    )

    heading2(doc, '6.6 Network Score Card (Kartu Skor Kesehatan)')
    body(doc,
        'Kartu ini mengkonversi data teknis yang kompleks menjadi satu angka (0-100) '
        'dan satu huruf grade (A/B/C/D) yang langsung dipahami:'
    )
    two_col(doc, [
        ('Grade A (Hijau, 90-100)', 'Sistem berjalan optimal — tidak ada anomali, CPU rendah'),
        ('Grade B (Biru, 75-89)',   'Sistem cukup baik — CPU agak tinggi tapi tidak anomali'),
        ('Grade C (Amber, 50-74)', 'Perhatian diperlukan — CPU tinggi atau mendekati anomali'),
        ('Grade D (Merah, <50)',   'Masalah aktif — anomali terdeteksi, perlu tindakan segera'),
    ])
    body(doc,
        'Saat anomali aktif, kartu ini berkedip-kedip (pulsing animation) untuk '
        'menarik perhatian operator.',
        italic=True
    )

    heading2(doc, '6.7 Panel Metode Deteksi (Model Stats Live)')
    body(doc,
        'Panel ini menampilkan "dapur" dari algoritma deteksi anomali secara transparan — '
        'sangat berguna saat presentasi untuk membuktikan bahwa sistem benar-benar '
        'menggunakan perhitungan statistik:'
    )
    bullet(doc, 'CPU mu/sigma: ',
           'Rata-rata dan standar deviasi beban CPU dari 100 log terakhir.')
    bullet(doc, 'Rx mu/sigma: ',
           'Rata-rata dan standar deviasi kecepatan download.')
    bullet(doc, 'Tx mu/sigma: ',
           'Rata-rata dan standar deviasi kecepatan upload.')
    bullet(doc, 'Z-Score aktif: ',
           'Nilai Z-Score real-time untuk setiap metrik. Jika Z > 3.0, '
           'angkanya berubah merah tebal.')

    heading2(doc, '6.8 Panel Client Terkoneksi')
    body(doc,
        'Menampilkan daftar perangkat WiFi yang saat ini terhubung ke router. '
        'Setiap perangkat ditampilkan dengan:'
    )
    bullet(doc, 'Nama Hostname: ',
           'Nama yang diberikan perangkat ke jaringan (misal: ASUS-ROG-Laptop, iPhone-15).')
    bullet(doc, 'IP Address: ',
           'Alamat IP lokal yang diberikan router (misal: 192.168.1.12).')
    bullet(doc, 'MAC Address: ',
           'Identifikasi fisik unik kartu jaringan perangkat — tidak bisa dipalsukan.')
    bullet(doc, 'SSID: ',
           'Nama WiFi yang digunakan perangkat (Winaila 5G atau Winaila 2.4G).')
    bullet(doc, 'Active Since: ',
           'Jam berapa perangkat pertama kali terdeteksi dalam sesi ini — '
           'berguna untuk identifikasi perangkat baru yang mencurigakan.')

    heading2(doc, '6.9 Log Database Historis')
    body(doc,
        'Menampilkan 12 log terbaru dari Supabase dalam bentuk card-card berwarna. '
        'Card berwarna merah muda = log anomali. Card abu-abu = log normal. '
        'Setiap card berisi: waktu, CPU%, kecepatan download, kecepatan upload, '
        'dan badge status NORMAL/ANOMALI.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # BAB 7 — CLOUDFLARE TUNNEL
    # ══════════════════════════════════════════
    heading1(doc, 'BAGIAN 7 — CLOUDFLARE TUNNEL: AKSES PUBLIK TANPA PORT FORWARDING')

    heading2(doc, '7.1 Masalah yang Dipecahkan')
    body(doc,
        'Secara normal, server yang berjalan di localhost:5000 hanya bisa diakses '
        'dari komputer yang sama. Untuk mengaksesnya dari luar (kampus, rumah teman), '
        'diperlukan "port forwarding" di router ISP — yang umumnya dilarang '
        'atau dikunci oleh IndiHome.'
    )

    heading2(doc, '7.2 Solusi: Cloudflare Tunnel')
    body(doc,
        'Cloudflare Tunnel (cloudflared) membuat terowongan aman dari server lokal '
        'ke infrastruktur Cloudflare, yang kemudian dihubungkan ke domain publik kita. '
        'Tidak ada port yang dibuka di router ISP — semua lalu lintas keluar '
        'melalui koneksi outbound yang aman.'
    )

    tip_box(doc,
        'Analogi Mudah untuk Audiens:',
        [
            '"Bayangkan server kita ada di dalam rumah yang tidak punya nomor alamat '
            'yang bisa ditemukan orang luar. Cloudflare Tunnel seperti menyewa '
            'kotak pos di kantor pos besar (Cloudflare) dengan alamat yang jelas '
            '(jarkom.clyuti.my.id). Semua surat yang datang ke kotak pos itu '
            'langsung diteruskan ke rumah kita secara otomatis."',
        ],
        bg='EFF6FF', border='93C5FD', title_color=(30,64,175)
    )

    two_col(doc, [
        ('URL Publik',        'https://jarkom.clyuti.my.id'),
        ('Port Lokal',        'localhost:5000 (backend + frontend)'),
        ('Protokol Tunnel',   'QUIC (UDP-based, lebih cepat dari TCP)'),
        ('Enkripsi',          'TLS 1.3 otomatis oleh Cloudflare'),
        ('Port Forwarding',   'TIDAK DIPERLUKAN — semua outbound'),
        ('Lokasi Server',     'sin02 (Singapore) & cgk05 (Jakarta) — 4 koneksi paralel'),
    ], headers=['Properti', 'Nilai'])

    heading2(doc, '7.3 Mengapa Bisa Diakses Dari WiFi Lain?')
    body(doc,
        'Pertanyaan yang sering muncul: "Jika ini monitoring jaringan kita, '
        'kenapa bisa diakses dari WiFi lain?" Jawabannya:'
    )
    bullet(doc, 'Yang Dipantau: ',
           'Traffic jaringan WiFi ZTE F670L di lokasi kita. Data diambil '
           'melalui scraping halaman admin router di 192.168.1.1.')
    bullet(doc, 'Yang Diakses: ',
           'Dashboard web di https://jarkom.clyuti.my.id — ini adalah '
           'ANTARMUKA yang menampilkan data, bukan router itu sendiri.')
    bullet(doc, 'Analoginya: ',
           'Seperti kamera CCTV di rumah yang bisa ditonton dari ponsel di luar. '
           'Kamera ada di rumah, rekaman bisa dilihat dari mana saja selama '
           'ada koneksi internet.')

    doc.add_page_break()

    # ══════════════════════════════════════════
    # BAB 8 — KETERBATASAN & PENANGANAN
    # ══════════════════════════════════════════
    heading1(doc, 'BAGIAN 8 — KETERBATASAN SISTEM & CARA MENGATASINYA')

    heading2(doc, '8.1 Tidak Bisa Packet Sniffing (Per-Device Bandwidth)')
    body(doc,
        'NetVision Pro tidak dapat mendeteksi SIAPA yang menggunakan bandwidth '
        'paling besar. Kita hanya bisa melihat total bandwidth seluruh jaringan, '
        'bukan penggunaan per-perangkat.'
    )
    body(doc, 'Mengapa tidak bisa?')
    bullet(doc, 'Posisi kita: ',
           'Komputer yang menjalankan backend hanya sebagai "client" di jaringan, '
           'bukan sebagai router atau gateway. Paket data antar perangkat lain '
           'tidak melewati network card kita.')
    bullet(doc, 'Router dikunci ISP: ',
           'ZTE F670L dari IndiHome tidak menyediakan akses CLI (SSH/Telnet) '
           'atau statistik per-device melalui halaman admin.')
    body(doc,
        'Solusi yang bisa dilakukan saat ini: Korelasikan waktu anomali dengan '
        'daftar client yang muncul di dashboard — perangkat yang baru muncul '
        'saat anomali terjadi kemungkinan besar adalah penyebabnya.',
        italic=True
    )

    heading2(doc, '8.2 Active Since Hilang Saat Backend Restart')
    body(doc,
        'Informasi "Active Since" (sejak kapan perangkat terhubung) hanya tersimpan '
        'di memori RAM backend. Setiap kali backend di-restart, waktu ini direset.'
    )
    body(doc,
        'Ini bisa diperbaiki dengan menyimpan data client_first_seen ke Supabase — '
        'namun untuk versi saat ini, ini adalah trade-off yang diterima.',
        italic=True
    )

    heading2(doc, '8.3 Hanya Berjalan di Windows')
    body(doc,
        'Pengambilan data CPU menggunakan perintah PowerShell (Get-CimInstance '
        'Win32_Processor) yang hanya tersedia di Windows. Untuk Linux/macOS, '
        'kode perlu dimodifikasi menggunakan perintah yang berbeda.'
    )

    heading2(doc, '8.4 Supabase Versi Gratis: Pause Setelah 7 Hari Tidak Aktif')
    body(doc,
        'Versi gratis Supabase akan meng-pause proyek jika tidak ada aktivitas '
        'selama 7 hari. Solusi: Pastikan backend berjalan setidaknya sekali '
        'seminggu, atau upgrade ke paket berbayar untuk proyek produksi.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # BAB 9 — Q&A PERSIAPAN DOSEN
    # ══════════════════════════════════════════
    heading1(doc, 'BAGIAN 9 — PERSIAPAN PERTANYAAN DARI DOSEN/PENGUJI')

    body(doc,
        'Berikut adalah pertanyaan-pertanyaan yang kemungkinan besar akan muncul '
        'saat presentasi, beserta jawaban yang direkomendasikan:'
    )

    qas = [
        ('Mengapa menggunakan web scraping, bukan SNMP?',
         'Router ZTE F670L yang disediakan IndiHome mengunci akses SNMP secara '
         'default. Firmware ISP menonaktifkan fitur ini untuk alasan keamanan '
         'dan kontrol jaringan. Web scraping melalui Puppeteer adalah solusi '
         'alternatif yang memanfaatkan halaman admin web yang memang disediakan '
         'oleh router. Teknik ini cukup handal selama struktur HTML router '
         'tidak berubah akibat firmware update.'),

        ('Apakah scraping ini legal/etis?',
         'Ya. Kita mengakses router kita sendiri menggunakan kredensial yang sah '
         '(username dan password milik kita sendiri). Ini bukan "hacking" — '
         'sama seperti kita masuk ke halaman admin router melalui browser biasa, '
         'hanya saja diautomatisasi menggunakan kode program.'),

        ('Bagaimana akurasi deteksi anomalinya?',
         'Algoritma Z-Score dengan threshold 3.0 mengikuti prinsip Three-Sigma Rule '
         'dalam statistik, yang menyatakan bahwa 99.7% data normal akan berada '
         'dalam rentang 3 standar deviasi dari rata-rata. Dengan kata lain, '
         'hanya 0.3% data normal yang berpotensi salah terdeteksi sebagai anomali. '
         'Ditambah dual-guard condition, akurasi praktisnya sangat tinggi.'),

        ('Kenapa ada data dari jam 11 padahal baru dijalankan jam 2?',
         'Data tersebut berasal dari sesi backend sebelumnya yang tersimpan permanen '
         'di Supabase Cloud. Ketika backend di-restart, data di database tidak ikut '
         'terhapus. Dashboard selalu menampilkan 50 log terbaru dari Supabase — '
         'yang mencakup semua sesi sebelumnya. Ini justru keunggulan sistem kami: '
         'riwayat tidak hilang meskipun backend di-restart.'),

        ('Apakah bisa digunakan untuk router merek lain?',
         'Untuk saat ini, implementasi scraping dikustomisasi khusus untuk antarmuka '
         'web ZTE F670L. Untuk router merek lain, perlu riset ulang untuk menemukan '
         'elemen HTML yang menyimpan data bandwidth. Arsitektur backend dirancang '
         'modular sehingga fungsi scrapeRouterStats() bisa diganti tanpa mengubah '
         'komponen lain.'),

        ('Seberapa sering data diperbarui?',
         'Ada dua frekuensi yang berbeda: Backend mengambil data dari router setiap '
         '30 detik dan menyimpannya ke Supabase. Frontend React mengambil data '
         'dari backend setiap 5 detik. Jadi tampilan dashboard diperbarui setiap '
         '5 detik, tapi nilai Rx/Tx yang berubah secara signifikan hanya setiap '
         '30 detik saat siklus scraping baru selesai.'),

        ('Bagaimana keamanan aplikasi ini?',
         'Saat ini dashboard dapat diakses siapa saja yang mengetahui URL. '
         'Untuk keamanan tambahan, bisa ditambahkan autentikasi (login/password) '
         'di level Cloudflare (Cloudflare Access) atau di level aplikasi. '
         'Data sensitif seperti password router tersimpan di file .env yang '
         'tidak pernah dikirim ke frontend.'),

        ('Apa rencana pengembangan ke depan?',
         'Beberapa pengembangan yang sudah diidentifikasi: (1) Notifikasi push '
         'melalui Telegram/email saat anomali pertama kali terdeteksi, '
         '(2) Autentikasi dashboard untuk keamanan, '
         '(3) Monitoring per-device menggunakan Raspberry Pi sebagai gateway, '
         '(4) Portabilitas ke Linux/macOS dengan mengganti perintah PowerShell, '
         '(5) Penyimpanan data client_first_seen ke Supabase agar Active Since '
         'tidak hilang saat restart.'),
    ]

    for q, a in qas:
        qa_box(doc, q, a)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # BAB 10 — PENUTUP
    # ══════════════════════════════════════════
    heading1(doc, 'BAGIAN 10 — PENUTUP & POIN KUNCI PRESENTASI')

    heading2(doc, '10.1 Tiga Poin yang Harus Selalu Diingat')
    tip_box(doc,
        'Poin Utama Presentasi (Selalu Kembalikan ke Sini):',
        [
            '1. ADAPTIF: Sistem belajar sendiri dari data historis setiap 5 menit.',
            '   Tidak ada yang perlu dikonfigurasi ulang secara manual.',
            '',
            '2. PERSISTEN: Supabase memastikan tidak ada data yang hilang,',
            '   bahkan setelah restart berkali-kali. Ini yang membuat sistem',
            '   benar-benar siap pakai jangka panjang.',
            '',
            '3. AKSESIBEL: Cloudflare Tunnel membuat dashboard dapat diakses',
            '   dari mana saja tanpa konfigurasi router yang rumit.',
        ],
        bg='F0FDF4', border='86EFAC', title_color=(5,122,85)
    )

    heading2(doc, '10.2 Kalimat Penutup Presentasi')
    tip_box(doc,
        'Script Penutup:',
        [
            '"NetVision Pro bukan sekadar dashboard — ini adalah sistem monitoring '
            'yang benar-benar HIDUP. Ia terus belajar dari pola jaringan Anda, '
            'menyimpan semua riwayat secara permanen di cloud, dan dapat diakses '
            'dari mana saja. Kami mengembangkan ini sebagai solusi nyata untuk '
            'keterbatasan hardware router ISP yang tidak menyediakan akses SNMP, '
            'membuktikan bahwa kreativitas dalam memilih pendekatan teknis '
            'bisa mengatasi batasan infrastruktur."',
        ],
        bg='EFF6FF', border='93C5FD', title_color=(30,64,175)
    )

    heading2(doc, '10.3 Teknologi yang Digunakan (Ringkasan)')
    two_col(doc, [
        ('Backend',     'Node.js + Express.js + Puppeteer (web scraping headless)'),
        ('Frontend',    'React.js 19 + Vite + Tailwind CSS + Recharts'),
        ('Database',    'Supabase Cloud (PostgreSQL)'),
        ('Tunneling',   'Cloudflare Tunnel (cloudflared)'),
        ('OS Layer',    'Windows PowerShell (CPU & Uptime)'),
        ('Router',      'ZTE F670L IndiHome (target monitoring)'),
        ('Algoritma',   'Z-Score Multi-Dimensi + Dynamic Re-training (5 menit)'),
        ('Hosting',     'Localhost + Cloudflare Edge Network (sin02/cgk05)'),
    ])

    divider(doc)
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_end.paragraph_format.space_before = Pt(16)
    r_end = p_end.add_run(
        f'NetVision Pro  |  Kelompok 6  |  Jaringan Komputer\n'
        f'Dokumen dibuat: {datetime.date.today().strftime("%d %B %Y")}'
    )
    r_end.italic = True; r_end.font.name = 'Arial'
    r_end.font.size = Pt(9)
    r_end.font.color.rgb = RGBColor(148,163,184)

    doc.save(OUTPUT_PATH)
    print(f'\nDokumen berhasil disimpan ke:\n   {OUTPUT_PATH}')

if __name__ == '__main__':
    build()
