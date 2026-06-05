import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)

def add_heading(text, level=1):
    h = doc.add_heading(text, level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
        if level == 1:
            run.font.size = Pt(12)
            run.font.bold = True
    return h

def add_paragraph(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(text)
    p.alignment = align
    return p

def add_image_placeholder(text_label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f" [ TEMPAT GAMBAR: {text_label} ] \n(Silakan Copy-Paste Gambar Anda di Sini)")
    run.font.bold = True
    run.font.color.rgb = docx.shared.RGBColor(255, 0, 0) # Red color for visibility

add_heading('HASIL DAN PEMBAHASAN', level=1)
add_paragraph('Bagian ini memaparkan hasil pengujian dari sistem NetVision Pro berdasarkan skenario metodologi yang telah dirancang. Pengujian berfokus pada keberhasilan pengumpulan data (scraping), akurasi deteksi anomali, serta fungsionalitas antarmuka dasbor.')

add_heading('1. Hasil Otomatisasi Scraping dan Sinkronisasi Cloud', level=2)
add_paragraph('Pengujian tahap pertama membuktikan bahwa backend berbasis Node.js dan Puppeteer berhasil melakukan login secara otomatis ke antarmuka router lokal (ZTE F670L). Sistem terbukti mampu mengekstrak metrik lalu lintas jaringan (akumulasi byte Rx dan Tx) serta mengonversinya menjadi kecepatan bandwidth aktual dalam satuan Mbps secara real-time. Data hasil ekstraksi tersebut berhasil disinkronisasikan ke basis data awan (cloud) Supabase secara konsisten setiap 60 detik. Gambar 1 menunjukkan cuplikan tabel network_logs pada dasbor Supabase yang merekam lebih dari 1.800 baris data historis secara kontinu tanpa ada kegagalan koneksi (timeout). Hal ini memvalidasi bahwa arsitektur pengumpulan data tanpa agen (agentless) yang diusulkan mampu bekerja dengan sangat stabil.')

add_image_placeholder("SCREENSHOT TABEL SUPABASE")
p = add_paragraph('Gambar 1. Log Data pada Basis Data Supabase', WD_ALIGN_PARAGRAPH.CENTER)
p.runs[0].font.bold = True

add_heading('2. Evaluasi Deteksi Anomali Z-Score Dinamis', level=2)
add_paragraph('Nilai tambah utama dari sistem ini adalah kemampuannya dalam mendeteksi aktivitas jaringan yang tidak wajar secara otomatis menggunakan algoritma statistik Z-Score. Pengujian membuktikan bahwa mekanisme latihan dinamis (dynamic retraining) berfungsi dengan sangat baik. Tabel 1 memaparkan sampel log saat anomali terdeteksi oleh sistem pada pengujian awal bulan Juni.')

# Add Table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Waktu Pengujian (WIB)'
hdr_cells[1].text = 'Parameter Metrik'
hdr_cells[2].text = 'Nilai Aktual'
hdr_cells[3].text = 'Keterangan Deteksi Anomali'

data = [
    ('3 Juni 2026, 00:14', 'Upload Speed', '21.12 Mbps', 'Upload abnormal (rata-rata normal ~4 Mbps)'),
    ('3 Juni 2026, 00:46', 'CPU Load', '97%', 'CPU Load abnormal (rata-rata normal ~41%)'),
    ('3 Juni 2026, 00:56', 'CPU Load', '84%', 'CPU Load abnormal (rata-rata normal ~42%)'),
    ('3 Juni 2026, 01:07', 'CPU Load', '99%', 'CPU Load abnormal (rata-rata normal ~41%)'),
]
for item in data:
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]
    row_cells[3].text = item[3]

p = add_paragraph('Tabel 1. Sampel Log Deteksi Anomali Jaringan', WD_ALIGN_PARAGRAPH.CENTER)
p.runs[0].font.bold = True

add_paragraph('Berdasarkan Tabel 1, terlihat jelas bahwa sistem secara adaptif menyesuaikan nilai "normal" (rata-rata historis/μ) seiring berjalannya waktu, dari 41% berubah ke 42%. Saat nilai aktual secara tiba-tiba melampaui batas ambang batas 3 kali standar deviasi (Z > 3), sistem langsung memicu status anomali. Sistem peringatan dini (early warning system) yang terintegrasi dengan fungsi pencatatan tersebut juga terbukti berjalan sukses dengan meneruskan notifikasi secara instan melalui bot Telegram. Gambar 2 menampilkan tangkapan layar peringatan yang diterima pengguna, yang memuat rincian mendalam (waktu, jenis anomali, dan deviasi nilai) sesaat setelah aktivitas anomali terjadi.')

add_image_placeholder("SCREENSHOT NOTIFIKASI TELEGRAM")
p = add_paragraph('Gambar 2. Notifikasi Deteksi Anomali melalui Telegram', WD_ALIGN_PARAGRAPH.CENTER)
p.runs[0].font.bold = True

add_heading('3. Implementasi Antarmuka Dasbor dan Aksesibilitas Publik', level=2)
add_paragraph('Hasil akhir dari pengolahan data sistem divisualisasikan melalui dasbor berbasis web. Gambar 3 memperlihatkan antarmuka utama NetVision Pro yang dirancang secara ergonomis menggunakan skema tata letak Bento Grid. Visualisasi grafik garis (line chart) lalu lintas jaringan secara interaktif menampilkan kurva pergerakan metrik unduh (download) dan unggah (upload). Sesuai dengan rancangan awal, sistem berhasil memberikan penanda visual diagnostik yang sangat intuitif dengan membubuhkan "Titik Anomali" berwarna merah menyala persis pada koordinat kurva di mana anomali statistik Z-Score sedang terjadi.')

add_image_placeholder("SCREENSHOT DASHBOARD BAGIAN ATAS (GRAFIK UTAMA)")
p = add_paragraph('Gambar 3. Antarmuka Utama Dasbor NetVision Pro', WD_ALIGN_PARAGRAPH.CENTER)
p.runs[0].font.bold = True

add_paragraph('Selain grafik utama, antarmuka pada Gambar 4 menunjukkan transparansi algoritma dengan menampilkan rincian parameter matematis Z-Score (μ/σ dan ambang batas deteksi) secara langsung di panel "Metode Deteksi". Pengujian aksesibilitas jaringan juga menunjukan hasil yang memuaskan. Dasbor yang secara fisik berjalan di server lokal (localhost) telah sukses diekspos agar dapat diakses dari jaringan publik eksternal menggunakan URL http://jarkom.clyuti.my.id. Keberhasilan akses URL publik ini didukung penuh oleh konfigurasi arsitektur Cloudflare Tunnel yang mampu mem-bypass perlindungan Network Address Translation (NAT) secara aman tanpa mengharuskan pengguna melakukan konfigurasi port forwarding pada router ISP.')

add_image_placeholder("SCREENSHOT DASHBOARD BAGIAN BAWAH (METODE Z-SCORE)")
p = add_paragraph('Gambar 4. Tampilan Detail Log Historis dan Parameter Algoritma Z-Score', WD_ALIGN_PARAGRAPH.CENTER)
p.runs[0].font.bold = True

add_heading('KESIMPULAN', level=1)
add_paragraph('Berdasarkan hasil perancangan dan pengujian yang telah dilakukan, dapat disimpulkan bahwa sistem pemantauan jaringan lokal NetVision Pro berhasil dibangun dan beroperasi secara terintegrasi dengan baik. Pemanfaatan teknik web scraping menggunakan Puppeteer terbukti menjadi solusi otomatisasi alternatif yang tangguh untuk mengekstrak data telemetri dari antarmuka web router (ZTE F670L) yang bersifat tertutup dan tidak menyediakan integrasi API publik. Implementasi algoritma statistik Z-Score multi-dimensi dengan mekanisme latihan dinamis (dynamic retraining) sukses mendeteksi keberadaan pola anomali—baik pada fluktuasi bandwidth maupun beban CPU—secara presisi tinggi dibandingkan dengan penggunaan ambang batas statis. Akurasinya semakin dikukuhkan melalui kecepatan respons penyampaian peringatan dini via bot Telegram. Pada sisi ujung depan (frontend), pemanfaatan pustaka React.js berhasil menghadirkan visualisasi dasbor Bento Grid yang interaktif lengkap dengan penanda titik merah pada riwayat anomali. Dilengkapi dengan infrastruktur Cloudflare Tunnel, sistem ini memberikan tingkat aksesibilitas yang sangat fleksibel, memungkinkan pengguna untuk memantau kesehatan jaringannya secara mandiri, real-time, dan aman dari manapun berada.')

doc.save('Draft_Jurnal_Siap_Copas.docx')
print("Document saved successfully.")
