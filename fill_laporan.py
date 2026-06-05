import docx
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_para_after(ref_element, text, bold=False, italic=False):
    """Insert a new paragraph element after ref_element in the document body."""
    new_p = OxmlElement('w:p')
    if text:
        pPr = OxmlElement('w:pPr')
        new_p.append(pPr)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        if italic:
            i_elem = OxmlElement('w:i')
            rPr.append(i_elem)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        new_p.append(r)
    ref_element.addnext(new_p)
    return new_p

def fill_report():
    doc_path = r"d:\jarkom\Laporan Progress_Kelompok 6.docx"
    doc = docx.Document(doc_path)
    
    # Find headings
    para_gambaran = None
    para_progress = None
    para_bukti = None
    
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt == 'Gambaran Sistem':
            para_gambaran = para
        elif txt == 'Progress Pengerjaan Projek':
            para_progress = para
        elif txt == 'Bukti Pengerjaan Projek':
            para_bukti = para
    
    if not (para_gambaran and para_progress and para_bukti):
        print("ERROR: Tidak dapat menemukan semua heading!")
        return
    
    print("✅ Semua heading ditemukan. Mulai mengisi konten...")
    
    # ============================================================
    # SECTION: GAMBARAN SISTEM
    # ============================================================
    gambaran_texts = [
        ("Sistem Monitoring Jaringan Lokal Berbasis Web (NetVision Pro) ini dirancang dengan arsitektur client-server terdistribusi untuk mengumpulkan metrik performa jaringan secara langsung dari perangkat router, memproses data tersebut di sisi backend, menyimpannya di database cloud, dan menampilkannya pada dashboard berbasis web secara real-time. Sistem ini juga dilengkapi dengan model deteksi anomali berbasis kecerdasan statistik dinamis menggunakan metode Z-Score serta fitur pelacakan perangkat klien WiFi yang terhubung.", False),
        ("", False),
        ("A. Arsitektur dan Komponen Sistem", True),
        ("", False),
        ("Sistem monitoring ini terdiri dari empat komponen utama yang saling berintegrasi:", False),
        ("", False),
        ("1. Data Collector (Backend Node.js + Puppeteer)", True),
        ("Bertindak sebagai mesin pengambil data otomatis (robot). Secara terjadwal setiap 60 detik, backend menjalankan browser headless (Chromium tanpa tampilan) menggunakan library Puppeteer untuk masuk ke halaman admin router ZTE F670L di alamat 192.168.1.1. Backend melakukan login otomatis, menavigasi ke menu Local Network > Status > WLAN Status, dan membaca akumulasi total byte yang diterima (download) dan dikirim (upload) dari elemen HTML TotalBytesCount:0. Selain itu, backend juga mengekspansi panel WLAN Client Status untuk membaca daftar perangkat yang sedang terhubung ke WiFi (hostname, IP, MAC Address, SSID). Backend juga mengambil data beban CPU dan uptime sistem operasi Windows melalui perintah PowerShell.", False),
        ("", False),
        ("2. Database Cloud (Supabase PostgreSQL)", True),
        ("Berfungsi sebagai penyimpan data log historis secara permanen di cloud. Setiap siklus polling, data terstruktur yang telah diproses dikirimkan dari backend ke tabel network_logs di Supabase Cloud menggunakan Supabase JS Client. Tabel ini menyimpan kolom: router_ip, cpu_usage, uptime, rx_bytes, tx_bytes, is_anomaly, dan anomaly_reason. Kredensial akses diamankan menggunakan variabel lingkungan (.env).", False),
        ("", False),
        ("3. Dashboard Pemantau (Frontend React.js + Recharts + Tailwind CSS)", True),
        ("Menyajikan visualisasi data yang interaktif dengan desain Bento Grid modern dan efek glassmorphism. Frontend melakukan polling data dari backend setiap 10 detik dan menampilkan: kecepatan unduh/unggah aktif (Mbps), persentase beban CPU dalam bentuk gauge chart setengah lingkaran, grafik tren lalu lintas data (line chart) dengan penanda titik anomali berwarna merah, grafik historis CPU Load, kartu Network Score (nilai 0-100 dan grade A/B/C/D), daftar perangkat klien WiFi yang terhubung beserta informasi Active Since, serta tabel riwayat log dari database.", False),
        ("", False),
        ("4. Tunnelling Publik (Cloudflare Tunnel)", True),
        ("Digunakan untuk mengekspos port backend dan frontend lokal ke internet publik melalui domain jarkom.clyuti.my.id. Hal ini memungkinkan dashboard dapat diakses dari mana saja secara real-time tanpa perlu konfigurasi port forwarding pada router ISP, sangat berguna saat presentasi di kampus.", False),
        ("", False),
        ("B. Mekanisme Pengumpulan Data (Web Scraping)", True),
        ("", False),
        ("Karena router ZTE F670L yang disediakan ISP tidak menyediakan API standar atau akses SNMP yang terbuka (dikunci oleh firmware ISP), proyek ini mengimplementasikan pendekatan web scraping menggunakan Puppeteer headless browser. Alur scraping yang berhasil diterapkan adalah sebagai berikut:", False),
        ("", False),
        ("a) Membuka halaman login router di http://192.168.1.1 menggunakan browser Chromium headless.", False),
        ("b) Memasukkan username dan password router pada form login (selector: #Frm_Username, #Frm_Password) lalu mengeklik tombol submit (#LoginId).", False),
        ("c) Navigasi ke menu utama Local Network (#localnet), dilanjutkan ke sub-menu Status (#localNetStatus).", False),
        ("d) Klik header bar WLAN Status (#WLANStatusBar) untuk memicu router memuat data bandwidth terbaru melalui mekanisme AJAX internal.", False),
        ("e) Klik header bar WLAN Client Status (#Wlan_ClientStatBar) untuk memuat daftar perangkat WiFi yang sedang terhubung.", False),
        ("f) Mengekstrak data total bytes received dan bytes sent dari elemen HTML dengan ID TotalBytesCount:0 yang memiliki format string 'rxBytes/txBytes' (contoh: '1262835434/3312516106').", False),
        ("g) Mengekstrak daftar klien WiFi dari elemen template_Wlan_ClientStat_N yang berisi hostname, IP Address, MAC Address, dan nama SSID.", False),
        ("h) Menghitung kecepatan unduh (Rx) dan unggah (Tx) aktif dalam satuan Mbps dengan rumus: Speed = (Delta_Bytes x 8) / (Delta_t x 10^6), di mana Delta_Bytes adalah selisih byte dari polling sebelumnya dan Delta_t adalah selang waktu dalam detik.", False),
        ("", False),
        ("C. Algoritma Deteksi Anomali Jaringan (Z-Score)", True),
        ("", False),
        ("Untuk memberikan nilai tambah berupa kecerdasan pada sistem pemantauan, kami mengimplementasikan deteksi anomali lalu lintas data dan beban CPU berbasis statistik Z-Score multi-dimensi:", False),
        ("", False),
        ("1) Proses Training Dinamis (Self-Learning): Setiap 5 menit sekali, backend mengambil 100 data log terbaru dari database Supabase untuk menghitung nilai Rata-rata (mean) dan Standar Deviasi (stdDev) dari tiga metrik utama: cpu_usage, rx_speed_mbps, dan tx_speed_mbps. Proses ini memungkinkan model beradaptasi secara otomatis terhadap pola penggunaan jaringan yang berubah-ubah.", False),
        ("", False),
        ("2) Perhitungan Z-Score: Setiap data baru masuk dari hasil polling, deviasi nilainya diukur dengan rumus: Z = |x - mean| / stdDev. Jika skor Z melampaui ambang batas sensitivitas (ANOMALY_THRESHOLD = 3.0) DAN metrik melampaui ambang batas absolut minimum (CPU > 60%, Rx > 50 Mbps, Tx > 20 Mbps), sistem menandai data tersebut sebagai anomali dan mencatat alasan spesifiknya.", False),
        ("", False),
        ("3) Mekanisme Cold Start (Fallback): Jika database memiliki kurang dari 20 record sehingga standar deviasi belum stabil untuk perhitungan statistik yang akurat, sistem secara otomatis beralih ke batas statis konservatif (CPU > 85%, Rx > 150 Mbps, Tx > 50 Mbps) sebagai pengaman sementara hingga data cukup terkumpul.", False),
        ("", False),
        ("D. Pelacakan Perangkat Klien WiFi (Active Since)", True),
        ("", False),
        ("Backend memiliki memori internal berupa Map global (clientFirstSeen) yang mencatat waktu pertama kali sebuah MAC Address terdeteksi terhubung ke jaringan WiFi. Setiap siklus polling, jika ditemukan MAC Address baru yang belum pernah tercatat, sistem akan menyimpan timestamp saat itu sebagai waktu Active Since. Informasi ini ditampilkan di dashboard sehingga pengguna dapat mengetahui sudah berapa lama setiap perangkat terhubung ke jaringan.", False),
        ("", False),
        ("E. Network Score (Penilaian Kesehatan Jaringan)", True),
        ("", False),
        ("Frontend menghitung skor kesehatan jaringan secara real-time (0-100 poin) berdasarkan kondisi saat ini. Skor dimulai dari 100 dan dikurangi berdasarkan: beban CPU tinggi (>50% dikurangi 5 poin, >70% dikurangi 15 poin, >90% dikurangi 30 poin) dan status anomali aktif (dikurangi 40 poin). Skor kemudian dikonversi menjadi grade huruf: A (>=90), B (>=75), C (>=50), D (<50). Kartu Network Score berubah warna secara dinamis sesuai grade.", False),
        ("", False),
        ("F. Teknologi yang Digunakan", True),
        ("", False),
        ("Backend: Node.js, Express.js, Puppeteer (headless Chrome), Supabase JS Client, dotenv, cors.", False),
        ("Frontend: React.js 19, Vite 8, Tailwind CSS 3, Recharts, Axios, Lucide React Icons.", False),
        ("Database: Supabase Cloud (PostgreSQL).", False),
        ("Infrastruktur: Cloudflare Tunnel (cloudflared), Windows PowerShell.", False),
        ("Router Target: ZTE F670L (firmware ISP IndiHome).", False),
    ]
    
    # Insert in REVERSE order after para_gambaran so they appear in correct order
    last_elem = para_gambaran._element
    for text, bold in gambaran_texts:
        last_elem = add_para_after(last_elem, text, bold=bold)
    
    print("  ✅ Bagian 'Gambaran Sistem' telah diisi.")
    
    # ============================================================
    # SECTION: PROGRESS PENGERJAAN PROJEK
    # ============================================================
    intro_text = "Proyek pengembangan dashboard monitoring jaringan ini telah diselesaikan hingga tahap akhir dengan fungsionalitas penuh. Seluruh komponen backend, frontend, database cloud, dan integrasi hardware telah berhasil diimplementasikan dan diuji. Berikut adalah ringkasan progress pengerjaan proyek Kelompok 6:"
    
    last_elem = para_progress._element
    last_elem = add_para_after(last_elem, intro_text)
    spacer = add_para_after(last_elem, "")
    
    # Create progress table
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Header row
    headers = ['No.', 'Tahap Pengerjaan', 'Detail Implementasi', 'Status']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_background(hdr_cells[i], '1F4E79')
        set_cell_margins(hdr_cells[i])
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
    
    # Data rows
    tasks = [
        ("1", "Setup Cloud Database", 
         "Inisialisasi project Supabase Cloud, pembuatan tabel network_logs dengan kolom: id, router_ip, cpu_usage, uptime, rx_bytes, tx_bytes, is_anomaly, anomaly_reason, created_at. Konfigurasi kebijakan Row-Level Security (RLS) agar backend dapat melakukan operasi INSERT dan SELECT secara aman.",
         "Selesai"),
        ("2", "Backend Core Development", 
         "Pembangunan server Express.js dengan endpoint REST API (/api/metrics/current, /api/metrics/history, /api/health). Implementasi pengambilan metrik CPU dan Uptime sistem Windows menggunakan perintah PowerShell. Integrasi Supabase JS Client untuk sinkronisasi data log berkala.",
         "Selesai"),
        ("3", "Router Web Scraping", 
         "Riset dan reverse-engineering halaman admin router ZTE F670L untuk menemukan elemen data bandwidth dan daftar klien WiFi. Implementasi Puppeteer headless browser untuk login otomatis, navigasi menu (Local Network > Status > WLAN Status + WLAN Client Status), ekspansi panel AJAX, parsing data byte dari TotalBytesCount:0, dan ekstraksi daftar perangkat terhubung. Penanganan error session dan browser crash recovery.",
         "Selesai"),
        ("4", "Model Deteksi Anomali", 
         "Pengembangan algoritma Z-Score multi-dimensi dengan mekanisme self-training dinamis setiap 5 menit menggunakan 100 data historis terakhir dari Supabase. Implementasi fallback batas statis untuk kondisi cold-start (data < 20 record). Konfigurasi threshold sensitivitas (Z > 3.0) dengan ambang absolut minimum.",
         "Selesai"),
        ("5", "Frontend Dashboard UI", 
         "Perancangan dan pembangunan antarmuka dashboard modern menggunakan React.js 19 dengan layout Bento Grid dan efek glassmorphism (Tailwind CSS). Implementasi komponen: kartu kecepatan download/upload, gauge chart CPU, grafik live traffic dengan penanda titik anomali merah, grafik CPU Load History, kartu Network Score (grade A-D), daftar client WiFi terkoneksi dengan Active Since, dan tabel log historis Supabase.",
         "Selesai"),
        ("6", "Tunnelling & Mock Mode", 
         "Konfigurasi Cloudflare Tunnel (cloudflared) untuk pemetaan port lokal (5000 dan 5173) ke domain publik jarkom.clyuti.my.id. Implementasi fitur Mock Mode di backend yang menghasilkan data simulasi dengan anomali terjadwal setiap ~2 menit untuk keperluan demonstrasi dan presentasi tanpa koneksi ke router fisik.",
         "Selesai"),
    ]
    
    for num, stage, detail, status in tasks:
        row_cells = table.add_row().cells
        row_cells[0].text = num
        row_cells[1].text = stage
        row_cells[2].text = detail
        row_cells[3].text = status
        
        for i, cell in enumerate(row_cells):
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                if i == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                elif i == 1:
                    for run in paragraph.runs:
                        run.bold = True
                elif i == 3:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(34, 139, 34)
    
    # Move table to correct position (after spacer)
    spacer.addnext(table._tbl)
    
    print("  ✅ Bagian 'Progress Pengerjaan Projek' telah diisi dengan tabel.")
    
    # ============================================================
    # SECTION: BUKTI PENGERJAAN PROJEK
    # ============================================================
    bukti_texts = [
        ("Sebagai bukti pengerjaan proyek, berikut adalah tangkapan layar (screenshot) dari sistem monitoring jaringan yang telah berhasil diimplementasikan dan berjalan secara aktif:", False),
        ("", False),
        ("1. Tampilan Dashboard Utama (Keadaan Normal)", True),
        ("Dashboard menampilkan seluruh metrik jaringan dalam layout Bento Grid dengan latar belakang gradasi terang dan efek glassmorphism. Terlihat kartu kecepatan download (cyan) dan upload (fuchsia), gauge CPU setengah lingkaran, grafik live traffic, kartu Network Score berwarna hijau dengan grade A, serta indikator jumlah client WiFi yang terhubung di header.", False),
        ("", False),
        ("2. Tampilan Dashboard Saat Anomali Terdeteksi", True),
        ("Ketika sistem mendeteksi anomali (Z-Score > 3.0), kartu Network Score berubah menjadi merah dengan efek berkedip (pulsing animation) dan grade turun menjadi D. Pada grafik traffic, titik data yang mengalami anomali ditandai dengan dot merah terang yang juga berkedip. Alasan anomali ditampilkan secara spesifik pada kartu score.", False),
        ("", False),
        ("3. Grafik CPU Load History", True),
        ("Grafik garis terpisah yang menunjukkan tren beban prosesor (CPU) secara historis. Grafik ini membantu pengguna mengidentifikasi pola kapan saja komputer bekerja terlalu keras dan berkorelasi dengan kejadian anomali jaringan.", False),
        ("", False),
        ("4. Daftar Client WiFi Terkoneksi", True),
        ("Panel yang menampilkan daftar perangkat yang sedang terhubung ke jaringan WiFi router. Setiap perangkat ditampilkan dengan hostname, IP Address, MAC Address, nama SSID (2.4G/5G), dan informasi Active Since yang menunjukkan sejak jam berapa perangkat tersebut mulai terhubung.", False),
        ("", False),
        ("5. Database Log di Supabase Cloud", True),
        ("Tangkapan layar dari dashboard Supabase menunjukkan tabel network_logs yang berisi riwayat data monitoring yang tersimpan secara otomatis setiap 60 detik. Terlihat kolom-kolom data termasuk rx_bytes, tx_bytes, cpu_usage, is_anomaly, dan timestamp pencatatan.", False),
        ("", False),
        ("6. Halaman Router ZTE F670L yang Di-scrape", True),
        ("Tangkapan layar dari halaman admin router ZTE F670L pada menu Local Network > Status > WLAN Status yang menampilkan data TotalBytesCount dan daftar WLAN Client yang menjadi sumber data utama sistem monitoring. Data inilah yang diekstrak secara otomatis oleh Puppeteer setiap siklus polling.", False),
        ("", False),
        ("7. Backend Server Berjalan (Terminal Log)", True),
        ("Log terminal menunjukkan backend Node.js yang aktif melakukan polling setiap 60 detik, berhasil melakukan scraping router, menghitung kecepatan bandwidth, mendeteksi jumlah client, menjalankan training model Z-Score, dan menyimpan data ke Supabase Cloud secara berkala.", False),
        ("", False),
        ("8. Cloudflare Tunnel Aktif", True),
        ("Konfigurasi Cloudflare Tunnel yang memetakan port lokal ke domain publik jarkom.clyuti.my.id, memungkinkan akses dashboard dari jaringan eksternal tanpa port forwarding.", False),
    ]
    
    last_elem = para_bukti._element
    for text, bold in bukti_texts:
        last_elem = add_para_after(last_elem, text, bold=bold)
    
    print("  ✅ Bagian 'Bukti Pengerjaan Projek' telah diisi.")
    
    # ============================================================
    # SAVE
    # ============================================================
    output_path = r"d:\jarkom\Laporan_Progress_Kelompok6_Lengkap.docx"
    try:
        doc.save(output_path)
        print(f"\n🎉 Dokumen berhasil disimpan ke: {output_path}")
    except PermissionError:
        alt_path = r"d:\jarkom\Laporan_Lengkap_v2.docx"
        doc.save(alt_path)
        print(f"\n⚠️  File terkunci. Disimpan ke: {alt_path}")

if __name__ == "__main__":
    fill_report()
